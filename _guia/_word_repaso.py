# -*- coding: utf-8 -*-
"""EL ARTIFACT, EN WORD: lo mismo, en el mismo orden y con las mismas pantallas.

Lo pidio Francisco asi -- «un word exactamente igual al artifact» --, y «igual» aqui no puede
ser una tanda de pantallazos de la pagina: eso pesa el triple, no se busca, no se copia y en
un movil no hay quien lo lea. Igual es que lleve LO MISMO y EN EL MISMO ORDEN:

  · sus titulos, sus pies y sus explicaciones, como texto de verdad,
  · sus maquetas -- que son HTML y CSS de la app y no hay forma de rehacerlas en Word --
    como imagen, recortadas del artifact ya montado,
  · nuestro veredicto de cada punto como texto, con su frase por frase y el motivo de las
    que no se ven,
  · y nuestra captura de la app al lado.

Las piezas las saca `_guia/_word_extraer.js` del propio artifact, no de los datos sueltos:
asi el Word no puede decir una cosa distinta de la que dice la pagina.

Uso:  backend/venv/Scripts/python.exe _guia/_word_repaso.py
"""
import io
import json
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PIEZAS = os.path.join(RAIZ, "_guia", "_word_items.json")
FOTOS = os.path.join(RAIZ, "_guia", "_word_capturas")
SALIDA = os.path.join(os.path.expanduser("~"), "Desktop",
                      "12EN12 - Los tres documentos y Nutricion, punto por punto.docx")

TINTA = RGBColor(0x1A, 0x18, 0x14)
GRIS = RGBColor(0x6B, 0x64, 0x5A)
NARANJA = RGBColor(0xC9, 0x4A, 0x0C)
VERDE = RGBColor(0x2B, 0x66, 0x42)
ROJO = RGBColor(0xB0, 0x28, 0x1C)

#: El color del veredicto, por la clase que le puso el artifact.
COLOR_ESTADO = {"ok": VERDE, "matiz": NARANJA, "mal": ROJO,
                "despues": GRIS, "gris": GRIS}

ANCHO_PAGINA = 6.6      # pulgadas utiles con los margenes de abajo
ANCHO_NUESTRA = 2.9     # la captura de la app, a media caja


def sombrear(parrafo, color):
    """El fondo de un parrafo, para los recuadros de codigo y los veredictos."""
    pr = parrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pr.append(shd)


def borde_izq(parrafo, color):
    pr = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    izq = OxmlElement("w:left")
    izq.set(qn("w:val"), "single")
    izq.set(qn("w:sz"), "18")
    izq.set(qn("w:space"), "8")
    izq.set(qn("w:color"), color)
    bordes.append(izq)
    pr.append(bordes)


def p(doc, texto="", tam=10.5, negrita=False, cursiva=False, color=None,
      antes=0, despues=4, izq=0, mono=False, alinear=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(antes)
    par.paragraph_format.space_after = Pt(despues)
    if izq:
        par.paragraph_format.left_indent = Inches(izq)
    if alinear:
        par.alignment = alinear
    if texto:
        r = par.add_run(texto)
        r.bold = negrita
        r.italic = cursiva
        r.font.size = Pt(tam)
        r.font.color.rgb = color or TINTA
        if mono:
            r.font.name = "Consolas"
    return par


def imagen(doc, ruta, ancho):
    """La imagen, sin pasarse del ancho util y sin estirarla."""
    if not ruta or not os.path.exists(ruta):
        return
    with Image.open(ruta) as im:
        w, h = im.size
    # Una maqueta muy alta se encoge para que quepa en una pagina y no salga cortada.
    pulgadas = min(ancho, 9.0 * w / h) if h else ancho
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(6)
    par.add_run().add_picture(ruta, width=Inches(pulgadas))


def bloque_mono(doc, texto, color_fondo="F3F1ED"):
    for linea in (texto or "").split("\n"):
        par = p(doc, linea or " ", tam=8.5, mono=True, despues=0, izq=0.12)
        sombrear(par, color_fondo)
    p(doc, despues=6)


def veredicto(doc, it):
    color = COLOR_ESTADO.get(it.get("clase") or "", GRIS)
    hexc = {"ok": "2B6642", "matiz": "C94A0C", "mal": "B0281C"}.get(it.get("clase") or "", "6B645A")

    par = p(doc, it.get("estado") or "", tam=9, negrita=True, color=color, antes=6, despues=2, izq=0.1)
    borde_izq(par, hexc)
    if it.get("nota"):
        par = p(doc, it["nota"], tam=10, izq=0.1, despues=3)
        borde_izq(par, hexc)

    for f in it.get("frases") or []:
        marca = "✓  " if f["se_ve"] else "✗  "
        par = p(doc, marca + f["frase"], tam=9.5, izq=0.22, despues=1,
                color=TINTA if f["se_ve"] else GRIS)
        borde_izq(par, hexc)
        if f.get("porque"):
            par = p(doc, f["porque"], tam=8.5, cursiva=True, color=GRIS, izq=0.42, despues=2)
            borde_izq(par, hexc)

    for clave in ("donde", "forzada"):
        if it.get(clave):
            par = p(doc, it[clave], tam=8.5, cursiva=True, color=GRIS, izq=0.1, despues=2)
            borde_izq(par, hexc)

    if it.get("imagen"):
        imagen(doc, os.path.join(FOTOS, it["imagen"]), ANCHO_NUESTRA)
    p(doc, despues=4)


def main() -> None:
    items = json.load(io.open(PIEZAS, encoding="utf-8"))

    doc = Document()
    est = doc.styles["Normal"]
    est.font.name = "Calibri"
    est.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    # ── Portada ──
    p(doc, "12EN12", tam=11, negrita=True, color=NARANJA, despues=2)
    p(doc, "Los tres documentos y la revisión de Nutrición,\npunto por punto",
      tam=24, negrita=True, despues=8)
    p(doc, "Cada punto de los documentos de Jesús con la pantalla de la app al lado, y la "
           "revisión funcional de Nutrición con lo que se arregló. Es el mismo contenido y el "
           "mismo orden que el artifact.", tam=11, color=GRIS, despues=10)
    p(doc, "Las maquetas son suyas, recortadas del documento tal y como las entregó. Las "
           "capturas son de la app funcionando, y cuando hubo que forzar un estado para poder "
           "mirarlo se dice cuál y por qué.", tam=10, color=GRIS, despues=10)
    p(doc, "2 de septiembre de 2026", tam=10, negrita=True, despues=2)
    p(doc, "Todo lo arreglado está en producción.", tam=10, color=VERDE, despues=0)

    doc.add_page_break()

    primero = True
    for it in items:
        t = it["tipo"]

        if t == "documento":
            if not primero:
                doc.add_page_break()
            primero = False
            p(doc, it.get("eti") or "", tam=9, negrita=True, color=NARANJA, despues=2)
            p(doc, it.get("titulo") or "", tam=19, negrita=True, despues=6)
            for b in it.get("bajadas") or []:
                p(doc, b, tam=10, color=GRIS, despues=4)
            p(doc, despues=6)

        elif t == "bloque":
            p(doc, it["texto"].replace("\n", " · "), tam=12, negrita=True, color=NARANJA,
              antes=16, despues=4)

        elif t == "seccion":
            p(doc, it["texto"], tam=14, negrita=True, antes=14, despues=3)

        elif t == "sub":
            p(doc, it["texto"], tam=9.5, cursiva=True, color=GRIS, despues=6)

        elif t == "punto":
            p(doc, it["texto"], tam=12, negrita=True, antes=12, despues=3)
            if it.get("pie"):
                p(doc, it["pie"], tam=9.5, cursiva=True, color=GRIS, despues=4)
            if it.get("medida"):
                bloque_mono(doc, it["medida"])
            for linea in it.get("maqueta_texto") or []:
                p(doc, "· " + linea, tam=9.5, izq=0.2, despues=1)

        elif t == "maqueta":
            if it.get("rot"):
                p(doc, it["rot"], tam=10, negrita=True, antes=8, despues=2)
            imagen(doc, os.path.join(FOTOS, f'{it["foto"]}.png'), ANCHO_PAGINA)
            if it.get("pie"):
                p(doc, it["pie"], tam=9.5, cursiva=True, color=GRIS, despues=4)

        elif t in ("calendario", "lista"):
            bloque_mono(doc, it["texto"], color_fondo="EFEDE8")

        elif t == "menor":
            p(doc, "· " + it["texto"].replace("\n", " — "), tam=10, izq=0.15, despues=4)

        elif t == "pie":
            p(doc, it["texto"], tam=9.5, cursiva=True, color=GRIS, antes=8, despues=4)

        elif t == "veredicto":
            veredicto(doc, it)

    doc.save(SALIDA)
    print(f"{len(items)} piezas -> {SALIDA}")
    print(f"{os.path.getsize(SALIDA) / 1e6:.1f} MB")


if __name__ == "__main__":
    main()
