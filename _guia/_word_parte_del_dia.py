# -*- coding: utf-8 -*-
"""EL PARTE DEL DÍA, EN WORD: lo mismo que el artifact y en el mismo orden.

Se saca **del propio HTML del artifact**, no de los datos sueltos, por la razón de siempre
(ver `_guia/_word_repaso.py`): así el Word no puede decir una cosa distinta de la que dice la
página. Aquí no hay maquetas de Jesús que recortar, así que sale todo como texto de verdad:
se busca, se copia y se lee en un móvil.

Uso:  backend/venv/Scripts/python.exe _guia/_word_parte_del_dia.py
"""
import os

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FUENTE = os.path.join(RAIZ, "_internos_proceso", "parte_del_dia_0309.html")
SALIDA = os.path.join(os.path.expanduser("~"), "Desktop",
                      "12EN12 - El parte del 3 de septiembre.docx")

NARANJA = RGBColor(0xD4, 0x53, 0x1A)
GRIS = RGBColor(0x55, 0x4E, 0x40)
ROJO = RGBColor(0xA8, 0x3A, 0x2A)


def _texto_con_negritas(p, nodo, color=None):
    """Vuelca los hijos de `nodo` en el párrafo, respetando <b>, <code> y <em>.

    Word no tiene <b> dentro de un párrafo: hay que ir trozo a trozo. Y sin esto se pierde
    justo lo que hace legible el parte, que son las cifras y los nombres en negrita.
    """
    for hijo in nodo.children:
        if isinstance(hijo, NavigableString):
            t = str(hijo)
            if not t.strip():
                if t and not p.runs:
                    continue
                p.add_run(" ")
                continue
            # El espacio de delante también cuenta: `<b>Los decimales.</b> La misma comida»`
            # traía uno y `split()` se lo comía, así que salía «decimales.La misma».
            limpio = " ".join(t.split())
            if t[:1].isspace() and p.runs and not p.runs[-1].text.endswith((" ", "\n")):
                limpio = " " + limpio
            r = p.add_run(limpio + (" " if t.endswith((" ", "\n")) else ""))
            if color:
                r.font.color.rgb = color
            continue
        nombre = hijo.name
        if nombre == "br":
            p.add_run("\n")
            continue
        # UN ESPACIO ENTRE TROZOS SI EL HTML LO TENÍA. En la página el hueco lo da el salto
        # de línea del fuente; al pegar run tras run se perdía y salían «Inicio0acbf4f» o
        # «Los decimales.La misma comida».
        texto = " ".join(hijo.get_text().split())
        if p.runs and not p.runs[-1].text.endswith((" ", "\n", "«", "(")) and texto:
            p.add_run(" ")
        r = p.add_run(texto)
        if nombre in ("b", "strong"):
            r.bold = True
        if nombre in ("em", "i"):
            r.italic = True
        if nombre == "code":
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        if nombre == "span" and "marca-sha" in (hijo.get("class") or []):
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRIS
        elif color:
            r.font.color.rgb = color


def main() -> None:
    sopa = BeautifulSoup(open(FUENTE, encoding="utf-8").read(), "html.parser")
    hoja = sopa.select_one(".hoja")
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── La portada, con lo que dice la cabecera del artifact ──
    cab = hoja.find("header")
    s = doc.add_paragraph()
    r = s.add_run(cab.select_one(".sello").get_text(strip=True).upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = NARANJA

    t = doc.add_paragraph()
    r = t.add_run(cab.find("h1").get_text(strip=True))
    r.bold = True
    r.font.size = Pt(26)

    b = doc.add_paragraph()
    _texto_con_negritas(b, cab.select_one(".bajada"), color=GRIS)

    cifras = " · ".join(
        f"{d.find('b').get_text(strip=True)} {d.find('span').get_text(strip=True).lower()}"
        for d in cab.select(".dato"))
    c = doc.add_paragraph()
    r = c.add_run(cifras)
    r.bold = True
    r.font.color.rgb = NARANJA
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    piezas = 0
    for nodo in hoja.find_all(["h2", "p", "h3", "table", "ol", "div"], recursive=True):
        clases = nodo.get("class") or []

        if nodo.name == "h2":
            doc.add_page_break() if piezas else None
            p = doc.add_paragraph()
            r = p.add_run(nodo.get_text(strip=True).upper())
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = NARANJA
            piezas += 1

        elif nodo.name == "h3":
            # El sha del commit va dentro del h3; se conserva, que es la trazabilidad.
            p = doc.add_paragraph()
            _texto_con_negritas(p, nodo)
            for r in p.runs:
                if r.font.name != "Consolas":
                    r.bold = True
                    r.font.size = Pt(13)
            # EL RELOJ Y EL ESTADO VAN CON SU TÍTULO. En la página son pastillas al lado del
            # h3 y aquí se perdían: el Word se quedaba sin los tiempos, que es justo una de
            # las cosas que se pidieron. Se emiten en la línea de debajo.
            cab = nodo.parent if "doc-cab" in (nodo.parent.get("class") or []) else None
            if cab:
                # Con separador: el sha va DENTRO del «artifact ...» y sin el se pegaban.
                trozos = [x.get_text(" ", strip=True) for x in cab.select(".fecha, .reloj, .chip")]
                if trozos:
                    q = doc.add_paragraph()
                    r = q.add_run(" · ".join(trozos))
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = NARANJA
            piezas += 1

        elif nodo.name == "div" and "cab" in clases and "opcion" in (nodo.parent.get("class") or []):
            # La cabecera de una opción: la letra, el nombre y lo que cuesta.
            p = doc.add_paragraph()
            letra = nodo.select_one(".letra")
            nombre = nodo.select_one(".nombre")
            coste = nodo.select_one(".coste")
            r = p.add_run(f"{letra.get_text(strip=True)} · " if letra else "")
            r.bold = True
            r.font.color.rgb = NARANJA
            if nombre:
                r = p.add_run(nombre.get_text(strip=True))
                r.bold = True
            if coste:
                r = p.add_run(f"   [{coste.get_text(strip=True)}]")
                r.font.size = Pt(9.5)
                r.font.color.rgb = GRIS
            piezas += 1

        elif nodo.name == "p" and nodo.parent is not cab:
            if "pie" in clases:
                doc.add_paragraph()
            p = doc.add_paragraph()
            _texto_con_negritas(p, nodo, color=GRIS if {"intro", "antes", "pie"} & set(clases) else None)
            piezas += 1

        elif nodo.name == "table":
            filas = nodo.select("tbody tr")
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = "Light Grid Accent 2"
            for i, txt in enumerate(("Min", "Qué", "Hoy")):
                cel = tabla.rows[0].cells[i].paragraphs[0]
                r = cel.add_run(txt)
                r.bold = True
            for fila in filas:
                celdas = fila.find_all("td")
                nueva = tabla.add_row().cells
                nueva[0].paragraphs[0].add_run(celdas[0].get_text(strip=True))
                _texto_con_negritas(nueva[1].paragraphs[0], celdas[1])
                nueva[2].paragraphs[0].add_run(celdas[2].get_text(strip=True))
            piezas += 1

        elif nodo.name == "ol" and "decisiones" in clases:
            for i, li in enumerate(nodo.find_all("li"), 1):
                p = doc.add_paragraph()
                r = p.add_run(f"{i}. ")
                r.bold = True
                r.font.color.rgb = NARANJA
                _texto_con_negritas(p, li, color=GRIS)
            piezas += 1

        elif nodo.name == "div" and "bloqueo" in clases:
            # Los bloqueos van marcados: son lo que para el trabajo.
            p = doc.add_paragraph()
            r = p.add_run("BLOQUEO")
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = ROJO
            piezas += 1

    doc.save(SALIDA)
    print(f"{piezas} piezas -> {SALIDA}")


if __name__ == "__main__":
    main()
