"""
Pasa la guia de markdown a Word, con sus capturas dentro.

No hay pandoc en la maquina, asi que se hace a mano con python-docx. Cubre lo que usa el
documento: titulos, parrafos con negrita/cursiva/codigo, listas, tablas, imagenes, citas,
lineas separadoras y enlaces.

Uso:  venv/Scripts/python _guia/_a_word.py
"""
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AQUI = os.path.dirname(os.path.abspath(__file__))
NARANJA = RGBColor(0xFF, 0x67, 0x1F)
TINTA = RGBColor(0x1A, 0x1A, 0x2E)
GRIS = RGBColor(0x6B, 0x6B, 0x7B)

# La caja util de una A4 con los margenes que usamos.
ANCHO_UTIL_CM = 16.0
ALTO_UTIL_CM = 21.0


def _sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def _texto_con_formato(parrafo, texto):
    """Reparte el texto en trozos con negrita, cursiva, codigo y enlaces."""
    # Los enlaces se quedan solo con su etiqueta: en papel una URL larga estorba.
    texto = re.sub(r"\[([^\]]+)\]\((?:[^)]+)\)", r"\1", texto)

    partes = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            r = parrafo.add_run(parte[2:-2])
            r.bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            r = parrafo.add_run(parte[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            r = parrafo.add_run(parte[1:-1])
            r.italic = True
        else:
            parrafo.add_run(parte)


def _añadir_imagen(doc, ruta_rel, pie):
    ruta = os.path.join(AQUI, ruta_rel)
    if not os.path.exists(ruta):
        print(f"  AVISO: falta la imagen {ruta_rel}")
        return
    # Se escala por ancho, y si sale mas alta que la pagina, por alto.
    try:
        from PIL import Image
        with Image.open(ruta) as im:
            w, h = im.size
        ancho = ANCHO_UTIL_CM
        alto = ancho * h / w
        if alto > ALTO_UTIL_CM:
            alto = ALTO_UTIL_CM
            ancho = alto * w / h
    except Exception:
        ancho, alto = ANCHO_UTIL_CM, None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alto:
        p.add_run().add_picture(ruta, width=Cm(ancho), height=Cm(alto))
    else:
        p.add_run().add_picture(ruta, width=Cm(ancho))

    if pie:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(pie)
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = GRIS
    doc.add_paragraph()


def _tabla(doc, filas):
    """filas: lista de listas de celdas ya troceadas."""
    if not filas:
        return
    columnas = max(len(f) for f in filas)
    t = doc.add_table(rows=0, cols=columnas)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, fila in enumerate(filas):
        celdas = t.add_row().cells
        for j in range(columnas):
            texto = fila[j] if j < len(fila) else ""
            celda = celdas[j]
            celda.text = ""
            p = celda.paragraphs[0]
            _texto_con_formato(p, texto)
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                _sombrear(celda, "1A1A2E")
            elif i % 2 == 0:
                _sombrear(celda, "F7F7F9")
    doc.add_paragraph()


def convertir(md_path, docx_path, titulo, subtitulo):
    doc = Document()

    # Margenes y tipografia base
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TINTA
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\n12EN12")
    r.bold = True
    r.font.size = Pt(48)
    r.font.color.rgb = NARANJA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS
    doc.add_page_break()

    lineas = open(md_path, encoding="utf-8").read().split("\n")
    i = 0
    en_codigo = False
    buffer_tabla = []

    def volcar_tabla():
        nonlocal buffer_tabla
        if buffer_tabla:
            _tabla(doc, buffer_tabla)
            buffer_tabla = []

    while i < len(lineas):
        linea = lineas[i]
        stripped = linea.strip()

        # Bloques de codigo
        if stripped.startswith("```"):
            en_codigo = not en_codigo
            i += 1
            continue
        if en_codigo:
            p = doc.add_paragraph()
            r = p.add_run(linea)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Tablas
        if stripped.startswith("|") and stripped.endswith("|"):
            celdas = [c.strip() for c in stripped.strip("|").split("|")]
            # La fila de guiones que separa cabecera de cuerpo no se pinta.
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in celdas if c):
                buffer_tabla.append(celdas)
            i += 1
            continue
        volcar_tabla()

        if not stripped:
            i += 1
            continue

        # Separadores
        if re.fullmatch(r"-{3,}", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "DDDDDD")
            bdr.append(bottom)
            pPr.append(bdr)
            i += 1
            continue

        # Imagenes
        m = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            _añadir_imagen(doc, m.group(2), m.group(1))
            i += 1
            continue

        # Titulos
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            nivel = len(m.group(1))
            texto = re.sub(r"[#*`]", "", m.group(2)).strip()
            if nivel == 1:
                doc.add_page_break()
            h = doc.add_heading(level=min(nivel, 4))
            h.text = ""
            r = h.add_run(texto)
            r.font.color.rgb = NARANJA if nivel <= 2 else TINTA
            r.bold = True
            r.font.size = Pt({1: 20, 2: 15, 3: 12.5, 4: 11}[min(nivel, 4)])
            i += 1
            continue

        # Citas
        if stripped.startswith(">"):
            texto = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            _texto_con_formato(p, texto)
            for r in p.runs:
                r.font.color.rgb = GRIS
                r.italic = True
            i += 1
            continue

        # Listas
        m = re.match(r"^(\s*)[-*]\s+(.*)$", linea)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            if len(m.group(1)) >= 2:
                p.paragraph_format.left_indent = Cm(1.4)
            _texto_con_formato(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", linea)
        if m:
            p = doc.add_paragraph(style="List Number")
            _texto_con_formato(p, m.group(2))
            i += 1
            continue

        # Parrafo normal: se juntan las lineas seguidas
        trozos = [stripped]
        j = i + 1
        while j < len(lineas):
            sig = lineas[j].strip()
            if (not sig or sig.startswith(("#", "-", "*", ">", "|", "!", "```"))
                    or re.match(r"^\d+\.\s", sig)):
                break
            trozos.append(sig)
            j += 1
        p = doc.add_paragraph()
        _texto_con_formato(p, " ".join(trozos))
        i = j

    volcar_tabla()
    doc.save(docx_path)
    print(f"OK  {os.path.basename(docx_path)}")


if __name__ == "__main__":
    convertir(
        os.path.join(AQUI, "GUIA-COMPLETA-12EN12.md"),
        os.path.join(AQUI, "GUIA-COMPLETA-12EN12.docx"),
        "Guía completa de la aplicación",
        "Todas las pantallas, todas las funcionalidades y cómo se relacionan\n3 de agosto de 2026",
    )
    convertir(
        os.path.join(AQUI, "ANEXO-TECNICO-endpoints.md"),
        os.path.join(AQUI, "ANEXO-TECNICO-endpoints.docx"),
        "Anexo técnico",
        "Los 194 endpoints de la API, uno por uno\n3 de agosto de 2026",
    )
