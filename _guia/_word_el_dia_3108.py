# -*- coding: utf-8 -*-
"""EL WORD DE «El día»: cada punto del documento, con la prueba de la app al lado.

El documento de Jesus se queda como esta. Esto es lo otro: por cada cosa que pedia, que
pide, que hace hoy el sistema y UNA CAPTURA de la app de verdad, sacada con
`_guia/_pruebas_el_dia_3108.js`.

Uso:  backend/venv/Scripts/python.exe _guia/_word_el_dia_3108.py
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FOTOS = os.path.join(RAIZ, "_guia", "_pruebas_el_dia")
SALIDA = os.path.join(os.path.expanduser("~"), "Desktop", "12EN12 - El dia - punto por punto.docx")

NARANJA = RGBColor(0xC9, 0x4A, 0x0C)
GRIS = RGBColor(0x5A, 0x53, 0x4A)
VERDE = RGBColor(0x2B, 0x66, 0x42)


def sombrear(celda, color):
    tc = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:fill"), color)
    tc.append(sombra)


def titulo(doc, texto, tam=15, color=None, espacio_antes=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(tam)
    if color:
        r.font.color.rgb = color
    return p


def parrafo(doc, texto, tam=10, cursiva=False, color=None, antes=0, despues=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    r = p.add_run(texto)
    r.italic = cursiva
    r.font.size = Pt(tam)
    if color:
        r.font.color.rgb = color
    return p


def etiqueta(doc, texto, color=GRIS):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = color
    return p


def cita(doc, texto):
    """Lo que dice el documento, literal."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.rows[0].cells[0]
    sombrear(c, "FBF1E9")
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9.5)
    return t


def captura(doc, fichero, ancho=3.1, pie=None):
    # Un nombre suelto sale de las capturas de «El día»; con barra, de cualquier carpeta de
    # `_guia` (lo de la tarde vive en `_de_donde_bajo/`, `_generico/` y compañía).
    ruta = (os.path.join(RAIZ, "_guia", fichero) if "/" in fichero
            else os.path.join(FOTOS, fichero))
    if not os.path.exists(ruta):
        parrafo(doc, f"[falta la captura {fichero}]", color=NARANJA)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(ruta, width=Inches(ancho))
    if pie:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(10)
        r = q.add_run(pie)
        r.font.size = Pt(8)
        r.font.color.rgb = GRIS
        r.italic = True


# ── Los puntos, en el orden del documento ────────────────────────────────────
#
# Cada uno: donde vive en el documento, que pedia, que hace hoy y su captura.
PUNTOS = [
    ("A · EL CHECK-IN DEL DÍA", None, None, None, None, None),

    ("Todas a la vista, sin plegar nada",
     "El check-in del día › la línea bajo el título, y «Las nueve, todas a la vista»",
     "«Todas a la vista, sin plegar nada: nueve preguntas y las notas.»",
     "Se acabó el acordeón. Hasta el 31-08 iba una tarjeta encendida cada vez: se contestaba "
     "una, se plegaba con su resumen y se abría sola la siguiente. Ahora están las ocho "
     "abiertas, con sus estrellas y sus botones.\n\n"
     "Y con el acordeón se cayó el naranja de la tarjeta: marcaba «esta es la que toca», y "
     "con las ocho abiertas no distingue nada. En esta app el naranja quiere decir «te has "
     "pasado», así que ocho tarjetas naranjas se leen como ocho errores.",
     [("A0_pantalla_entera.png", 2.6, "La pantalla entera, con las ocho preguntas abiertas.")]),

    ("Fuera «Sensaciones generales del día»",
     "El check-in del día › «Las nueve, todas a la vista» (ya no está en la lista)",
     "La lista de las nueve ya no la incluye: era la primera pregunta y desaparece.",
     "Se va de la PANTALLA, no de la base. El campo se sigue guardando y el historial de "
     "quien lo tenga contestado lo sigue pintando con sus estrellitas: borrarlo del modelo "
     "dejaría meses de días sin poder enseñarse.\n\n"
     "La pantalla ahora abre directamente por «¿Entrenaste hoy?».",
     [("A1_primera_pregunta.png", 3.0, "La primera pregunta es ya «¿Entrenaste hoy?».")]),

    ("La suplementación entra en la cadena",
     "El check-in del día › «Las nueve, todas a la vista» · 5.ª línea",
     "«¿Tomaste la suplementación que tenías pautada? — Sí · No toda · No ← entra: es la que "
     "faltaba de tus once»",
     "La pregunta existía desde el 24-08 pero era CONDICIONAL: el servidor miraba que su plan "
     "incluyera suplementación y que tuviera protocolo vigente ese día, y solo entonces la "
     "pintaba. Por eso no salía en la columna «como está hoy» del documento: ese cliente no "
     "la tenía.\n\n"
     "Se le quita el candado y se le pregunta a todo el mundo, con el texto nuevo y con «No "
     "toda» en vez de «No todos».",
     [("A2_suplementacion.png", 3.0, "Con su texto nuevo y sus tres opciones.")]),

    ("El subtítulo de los extras",
     "El check-in del día › «Las nueve, todas a la vista» · 6.ª línea",
     "«¿Se te ha escapado algo más hoy? — Si no lo pusiste en el apartado de extras, ponlo "
     "ahora»",
     "Decía «Algo que comieras de más y no pusieras en el apartado “extras”», que describe. "
     "Esta pide. El campo ya llevaba al mismo sitio que los Extras del Inicio, y eso no se "
     "toca.",
     [("A3_extras.png", 3.0, "El subtítulo nuevo, debajo de la pregunta.")]),

    ("El «Te queda por contestar», entero",
     "El check-in del día › «Antes de guardar»",
     "«Lo único que cambia es el texto: sin sensaciones, que ya no existe, y sin el “y 5 "
     "más”. Las ocho enteras.»",
     "Cortaba en tres y añadía «y 5 más». Eso tenía sentido con el acordeón, porque la "
     "pantalla ya decía por dónde ibas; con las ocho abiertas, esta línea es lo único que lo "
     "dice, y un «y 5 más» obliga a releerse la pantalla para saber cuáles.",
     [("A4_pie.png", 3.1, "Las ocho, con la de la suplementación dentro.")]),

    ("El aviso de arriba, en sus dos estados",
     "El check-in del día › «Si le faltan comidas» y «Si no le falta nada»",
     "«Y el aviso de arriba, en sus dos estados. Hoy solo existe el de la izquierda: cuando "
     "no le falta nada, ese hueco se queda vacío.» · «Si “El día, todo bien” era un texto, "
     "este es su sitio: el mismo hueco, en verde.»",
     "El primero ya existía y se queda igual, palabra por palabra: «Te quedan 2 comidas sin "
     "registrar · Intra-entreno · Post-entreno · Puedes cerrarlas antes de seguir». Es "
     "también el trozo que el documento llama «Lo primero que ve» y pone entre los que no se "
     "tocan.\n\n"
     "El segundo es el nuevo. El hueco ya tenía un estado verde, pero decía «Dieta "
     "registrada» y solo salía si había dieta montada. Ese candado existía por algo: un día "
     "al que solo le apuntaron «dos cañas» tiene documento y ninguna comida, y decirle «dieta "
     "registrada» habría sido mentira. Con el texto nuevo el candado se cae, porque lo que se "
     "afirma ha cambiado: «no te queda nada por registrar» es verdad no haya comidas "
     "pendientes, haya montado dieta o no.",
     [("A5b_comidas_sin_registrar.png", 3.1,
       "Si le faltan comidas: montado un día con el intra y el post vacíos, la app los nombra."),
      ("A5_el_dia_todo_bien.png", 3.0,
       "Y si no le falta nada: el mismo hueco, en verde y con dos renglones.")]),

    ("Las notas, suyas o compartidas",
     "El check-in del día › «Las notas, suyas o compartidas»",
     "«Esto es para tu diario. Lo puedes compartir con nosotros o quedártelo para ti», con la "
     "casilla debajo. El cliente decide qué te enseña. Está bien resuelto.",
     "De los tres trozos que el documento marca como «no se tocan», este es el segundo. Se "
     "comprueba que sigue intacto: la frase, la casilla y el sitio, al final y abiertas.",
     [("A7_las_notas.png", 3.1, "Con su frase y su casilla, sin tocar.")]),

    ("Contestar ya no pliega la tarjeta",
     "El check-in del día › consecuencia de «todas a la vista»",
     "—",
     "Antes, contestar apagaba la tarjeta y encendía la siguiente. Ahora contestar solo "
     "guarda el valor: la tarjeta se queda abierta, con su tic, y se puede cambiar la "
     "respuesta sin volver a abrir nada.",
     [("A6_contestada_sigue_abierta.png", 3.0,
       "Contestada «No»: sigue abierta, con su tic y su respuesta marcada.")]),

    ("B · QUÉ PASA CUANDO DEJA DE CERRAR EL DÍA", None, None, None, None, None),

    ("Ese mismo día, desde las 17:00",
     "El check-in del día › «Ese mismo día, desde las 17:00»",
     "«Por la mañana no sale. Se enciende a las 17:00, hora de España: antes no tiene nada "
     "que cerrar, y verla apagada todo el día la convierte en parte del decorado.»",
     "La hora es la de España, no la del reloj del cliente: es la regla de la casa: su reloj "
     "decide qué día vive, España decide plazos y ventanas. Y es el mínimo, no un valor "
     "fijo: el cliente puede retrasarla desde su perfil (turnos de noche), nunca "
     "adelantarla.",
     [("B1_desde_su_hora.png", 3.1, "El estado normal, a partir de su hora.")]),

    ("Tras 2 días perdidos",
     "El check-in del día › «Tras 2 días perdidos»",
     "«¿Cómo fuiste hoy? / Llevas 2 días seguidos sin cerrar, no lo dejes hoy también». Y: "
     "«No lo dejes hoy también es lo que hace el trabajo: le pide el de hoy, no le riñe por "
     "los de atrás.»",
     "La racha cuenta SOLO los días que ya no tienen arreglo: por la mañana el de ayer sigue "
     "abierto, así que no cuenta todavía. Contarlo sería reñirle por algo que aún puede "
     "hacer.",
     [("B2_dos_dias.png", 3.1, "El título no cambia todavía; cambia lo de debajo.")]),

    ("Tras 4 días perdidos",
     "El check-in del día › «Tras 4 días perdidos»",
     "«Llevas 4 días sin cerrar / Retómalo hoy mismo: es de donde salen tus ajustes». Y: "
     "«Aquí se le dice lo que se pierde, y es verdad: cuando llegue su reporte, esos días van "
     "en blanco. Sin riña, que dos días antes le estabas animando.»",
     "Aquí cambia el título, no solo el subtítulo, y por primera vez se le dice lo que se "
     "está perdiendo.",
     [("B3_cuatro_dias.png", 3.1, "Cambia el título y dice qué se pierde.")]),

    ("Tras una semana",
     "El check-in del día › «Tras una semana»",
     "«Llevas una semana sin cerrar el día / Dejo de recordártelo. Si te está costando, "
     "dímelo y lo vemos». Y: «Tu frase, sin el “hasta que vuelvas”: la línea no desaparece. "
     "Si se quitara, se queda sin el único sitio donde se le dice y no vuelve.»",
     "A partir de la semana el texto se queda fijo: deja de escalar, pero la fila sigue ahí "
     "para siempre. A los veinte días dice lo mismo, no algo peor.\n\n"
     "Al sacar esta captura salió un fallo que no estaba en el documento: la fila recortaba "
     "el texto («Dejo de recordártelo. Si te está costando, dímelo y l…»), y esa frase es "
     "justo el trabajo. Se le quitó el recorte.",
     [("B4_una_semana.png", 3.1, "La frase entera, en tres renglones.")]),

    ("C · LA VENTANA DE LA MAÑANA", None, None, None, None, None),

    ("El día, con sus horas",
     "El día que no lo cierra › «El día, con sus horas»",
     "«17:00 — se abre el check-in de hoy. Al día siguiente, hasta las 15:00 — si ayer se "
     "quedó sin cerrar, sale el aviso y lo puede rellenar con la cabeza fresca. 15:00 — ayer "
     "se cierra. Ya no vuelve. 17:00 — se abre el de hoy. Nunca hay dos días abiertos a la "
     "vez, y entre las tres y las cinco hay dos horas de hueco justo para eso.»",
     "Es UNA SOLA VENTANA, como dice el documento: el cierre de un día está abierto desde su "
     "hora hasta las 15:00 del siguiente. De ahí sale todo lo demás, y el tope de arriba se "
     "resuelve solo.\n\n"
     "Esta captura es la franja de las 15:00 a las 17:00, con el servidor de verdad: no hay "
     "ningún día abierto, así que la fila del cierre no está en «Pendiente».",
     [("C3_sin_dia_abierto.png", 2.6,
       "Entre las 15:00 y las 17:00: en «Pendiente» no hay fila de cierre.")]),

    ("A la mañana siguiente, si ayer no lo cerró",
     "El día que no lo cierra › «A la mañana siguiente, si ayer no lo cerró»",
     "«Ayer no cerraste el día / Puedes hacerlo hasta las 3 de la tarde». Y: «Es el único "
     "sitio donde se recupera un día: al día siguiente, con la cabeza fresca. A las tres se "
     "cierra y ya no vuelve.»",
     "Es lo único del documento que no existía de ninguna forma. La fila lleva al día de "
     "AYER, no al de hoy.",
     [("C2_la_manana_siguiente.png", 3.1, "La fila de la mañana."),
      ("C2b_cierre_de_ayer.png", 3.1,
       "Y al tocarla se abre el día de ayer, no el de hoy.")]),

    ("D · LA CONFIGURACIÓN DE LOS AVISOS", None, None, None, None, None),

    ("Los siete interruptores y la hora",
     "La configuración de los avisos › «Todo encendido»",
     "Siete interruptores y un selector de hora, en cuatro grupos: el cierre del día, los "
     "reportes, el peso y cómo te aviso.",
     "Hasta hoy era UNO: «Recordarme cerrar el día · cada día a las 20:00».\n\n"
     "Y apagan de verdad, que era lo importante: «un interruptor que no hace nada enseña que "
     "la configuración miente».",
     [("D1_tarjeta.png", 3.0, "Los cuatro grupos, con sus tres frases de aviso.")]),

    ("La hora: una ventana, no dos cosas",
     "La configuración de los avisos › «La hora: una ventana, no dos cosas»",
     "«Puedes activarla a cualquier hora a partir de las 17:00. Permanecerá activa hasta las "
     "15:00 del día siguiente.» Y: «Tiene un motivo real y es tuyo: los turnos de noche. Al "
     "que sale a las dos de la mañana las 17:00 no le sirven.»",
     "El selector ofrece de las 17:00 a las 23:00 y ni una antes: no se puede cerrar un día "
     "que todavía no ha pasado. Y la frase de debajo es la que explica que la ventana de la "
     "mañana no es un mecanismo aparte, sino la misma.",
     [("D2_selector_de_hora.png", 3.1, "De 17:00 a 23:00, y la frase de la ventana.")]),

    ("Al apagar el cierre del día",
     "La configuración de los avisos › «Al apagar el cierre del día»",
     "«Si lo apagas, no podrás registrar tus datos del día, pero deberás rellenar las "
     "preguntas del reporte quincenal y del reporte mensual para poder recibir tus ajustes.» "
     "Y debajo: «Puedes volver a activarlo cuando quieras.»",
     "Se pregunta antes, con su texto literal y sus dos botones. La segunda línea no es "
     "adorno: «sin ella el interruptor parece una puerta de un solo sentido, y hay gente que "
     "no lo toca por miedo a no poder deshacerlo».\n\n"
     "Encenderlo no pregunta: volver atrás no le quita nada.",
     [("D4_dialogo_al_apagar.png", 3.1, "El diálogo, con «Dejarlo como está» y «Apagarlo».")]),

    ("Los dos del cierre del día son distintos",
     "La configuración de los avisos › «Los dos del cierre del día son distintos»",
     "«Rellenar el cierre del día apagado → la fila no sale nunca. Recordármelo si me lo "
     "salto apagado → la fila sale igual, pero no le insiste: la escalada de los 2, 4 y 7 "
     "días no salta. Hoy no puede elegir eso: o todo o nada.»",
     "Comprobado con el cliente de pruebas, que llevaba 5 días sin cerrar: con «Recordármelo» "
     "apagado, la fila vuelve a decir «Para rellenar al final del día» en vez de «Llevas 4 "
     "días sin cerrar». La escalada la apaga el servidor, no la pantalla.",
     [("D1_tarjeta.png", 3.0, "Los dos, separados, dentro del primer grupo.")]),

    ("E · EL PANEL", None, None, None, None, None),

    ("Una columna de días sin cerrar",
     "El check-in del día › «Los avisos no salen de la app», última línea",
     "«Lo que sí falta está en el panel: una columna de días sin cerrar en Clientes. Dejar de "
     "cerrar suele ser el primer síntoma, y llega antes que el impago.»",
     "La cuenta es la MISMA que la de la fila del Inicio, así que el panel y el cliente no "
     "pueden discrepar. En naranja a partir de cuatro días, que es cuando la app deja de "
     "animarle y le dice lo que se pierde.\n\n"
     "En el tope pone «+60» y no «60»: la cuenta se corta ahí y un número exacto que no es "
     "verdad es peor que uno redondo.",
     [("E1_columna_sin_cerrar.png", 6.2, "La columna «Sin cerrar», en Clientes.")]),

    ("Que se marque a quien lo apagó",
     "La configuración de los avisos › «Duda 7» y «Como te llega hoy»",
     "«Si no se marca, el que ignoró el reporte y el que tenía el aviso apagado te llegan "
     "iguales.» Y: «El que lo ignoró, en rojo. El que apagó los avisos, en gris y dicho. Su "
     "silencio no significa lo mismo.»",
     "Al que apagó el cierre no se le cuentan los días: pone «apagado». Contarlos sería el "
     "mismo error por el otro lado: no se ha dejado nada, es que no lo tiene.\n\n"
     "Y en reportes, el que lo ignoró va en naranja y el que apagó los avisos en gris, con "
     "«Avisos apagados».",
     [("E2_apagado_en_la_tabla.png", 6.2,
       "El mismo cliente con el cierre apagado: la columna lo dice en vez de contar.")]),

    # ── Lo de la tarde: no sale del documento de Jesús, sale de clientes usando la app ──
    ("F · LO QUE LLEGÓ DESPUÉS · LOS FALLOS QUE REPORTARON LOS CLIENTES", None, None, None, None, None),

    ("Copiar una favorita descuadraba los macros",
     "No está en el documento: lo reportó Gonzalo Rubio el 31-08",
     "«Al copiar un menú que ya guardé como favorito en otro día, se me desajustan los "
     "macros. Los que deberían ser: 3 comidas de 60 prote / 33,5 carbs / 16,5 grasa.»",
     "Al aplicar una favorita, la pantalla le pedía al servidor el reparto para CUATRO "
     "comidas siempre que la favorita tuviera tres. Era lo que quedaba del apaño de cuando el "
     "reparto no sabía de días de tres comidas: el 17-06 se quitó de los otros tres sitios y "
     "este se quedó.\n\n"
     "Medido contra el servidor: con 3 reparte 63,3 g de proteína por comida; con 4, "
     "47,5 / 47,5 / 38. Él veía 45 / 45 / 36 — distintos macros, el mismo 1,25 exacto entre "
     "la primera y la tercera. Esa es la huella. Y al día le faltaban los de una cuarta "
     "comida que no se pinta en ninguna parte: 57 g de proteína al aire.\n\n"
     "Arreglado y en producción desde el 31-08.",
     None),

    ("«Lentejas cocidas que se añadió solo»",
     "No está en el documento: lo reportó Gonzalo Rubio el 31-08",
     "«Cuando quiero copiar un post entreno que ya creé otro día, al copiarlo me aparecen "
     "diferentes cantidades y lo de lentejas cocidas, que obviamente se añadió solo.»",
     "Guardar un día NUNCA sustituye: fusiona. Eso es a propósito desde el 16-08, para que "
     "dos pestañas abiertas no se pisen. Pero de ahí salían dos cosas: una comida que la "
     "favorita no traía desaparecía de la pantalla y seguía guardada — volvía al recargar —, "
     "y «Copiar a otro día», que promete por escrito «la dieta de ese día se sustituye por "
     "esta», dejaba dentro lo que el origen no llevaba.\n\n"
     "Reproducido literal antes de tocar nada: se copió encima de un día con lentejas en la "
     "Comida 3 y ahí seguían. Ahora aplicar una favorita vacía lo que no trae, y copiar "
     "sustituye de verdad. Y se dice: cambiar un añadido silencioso por un borrado silencioso "
     "no arregla nada.",
     [("_favoritas_3108/3_tras_recargar.png", 3.0,
       "Tras aplicar la favorita y recargar: el intra con lentejas ya no vuelve.")]),

    ("Cuadrar pregunta de dónde bajar, en vez de decidirlo por el orden de la lista",
     "No está en el documento: nota de voz de Jesús del 31-08",
     "«Cuando recalcule, pregunte, pregunte de qué quiere bajar la proteína, del polvo o del "
     "queso. Es imposible que la aplicación aprenda eso, porque te puede quedar más denso, "
     "menos denso. Lo más sencillo es preguntar, o sea que pregunte de dónde recalcula.»",
     "Medido antes de tocar nada: una comida con 60 g de aislado y 300 g de queso batido "
     "para un objetivo de 38 g de proteína salía con el aislado en 5 g y el queso intacto. Y "
     "con los mismos dos alimentos, cambiando SOLO el orden de la lista, salía el queso en "
     "285 y el aislado en 10. La app ya decidía; lo que no tenía era criterio: mandaba el "
     "orden.\n\n"
     "Y no es cosa de batidos: pollo 250 g + atún 150 g + arroz 100 g salía con el pollo "
     "desplomado a 60 g y el atún casi entero en 125. Quien se hizo un plato de pollo con "
     "algo de atún se encontraba un plato de atún con algo de pollo.\n\n"
     "La pregunta solo dice lo que la app sabe: cuánto hay, cuánto pone del macro que sobra y "
     "en cuánto se quedaría. Nada de «queda más espeso»: la app no sabe la textura de nada, y "
     "en un plato de pollo con arroz esa frase no significaría nada. Ninguna opción viene "
     "marcada («ni siquiera que sugiera»).\n\n"
     "Si solo un alimento pone ese macro, no pregunta: lo baja y lo dice. Y mientras nadie "
     "conteste se baja en proporción, nunca por el orden de la lista.",
     [("_de_donde_bajo/1_la_pregunta.png", 3.4,
       "La pregunta, con las tres salidas y sus cantidades.")]),

    ("Aplicar una favorita no interroga: marca la comida",
     "No está en el documento: consecuencia de lo anterior",
     "—",
     "Una favorita recuadra cuatro o cinco comidas de golpe, y ahí un interrogatorio sería "
     "peor que el problema. Se baja en proporción, se avisa, y la comida queda marcada en su "
     "tarjeta con la pregunta a un toque.\n\n"
     "Para eso se guarda también lo que había ANTES: cuando el cliente toca la marca, las "
     "cantidades de ahora ya están bajadas y preguntar sobre ellas no daría ninguna opción.",
     [("_de_donde_bajo/4_la_marca.png", 5.2,
       "La marca en la tarjeta de la comida.")]),

    ("Y si bajar no llega, pregunta qué quitar — hasta que cuadre",
     "No está en el documento: lo probó Francisco el 31-08",
     "«Puse cuadrar, me dio a elegir entre 2 opciones, la resto pero sigue sin cuadrar. Para "
     "qué me dice que quite; también me debería preguntar qué quitar. Los macros tienen que "
     "quedar cuadrados, ese es el objetivo del botón.»",
     "Su comida tenía catorce alimentos y CON TODO A SU MÍNIMO PESABLE daba 57,9 P / 18,2 H / "
     "50,2 G contra un objetivo de 47,5 / 72 / 12: sobraban 38 g de grasa aunque no quedara "
     "nada que bajar. Esa comida no se puede cuadrar bajando, así que el modal no tenía "
     "ninguna salida que ofrecer y ni salía.\n\n"
     "Ahora la aritmética elige la forma de la pregunta: si con todo en el suelo el macro "
     "cabe, «¿de dónde bajo?»; si ni así cabe, «¿qué quito?». Y se sigue preguntando hasta "
     "que la comida cuadre o el cliente lo deje. No rompe la regla del 08-08 de que cuadrar "
     "no quita ingredientes: la app sigue sin quitar nada por su cuenta, lo quita él.\n\n"
     "De paso salió que el corte de quién entraba en la pregunta estaba al revés: era «que "
     "ponga al menos el 15 % del macro», y eso callaba la pregunta cuando un alimento "
     "dominaba — el bacon ponía la mitad de la grasa y era el único que pasaba, así que "
     "quedaba un solo candidato. Ahora el criterio es el margen: lo que pone menos lo que "
     "pondría en su mínimo.",
     [("_cuadra_hasta_cuadrar/2_como_queda.png", 6.2,
      "La comida de los catorce, ya cuadrada: 47,2 P y 72,9 H en verde.")]),

    ("Una sola pregunta, con la cuenta a la vista",
     "No está en el documento: lo pidió Francisco el 31-08",
     "«Va haciendo muchas veces las preguntas, supongo que lo hará hasta que cuadre pero es "
     "tedioso.»",
     "Era verdad: su comida pedía cinco diálogos seguidos, uno por alimento, y solo al cerrar "
     "el quinto se enteraba de que ya cuadraba. La decisión sigue siendo suya, pero cabía "
     "entera en una pantalla.\n\n"
     "Ahora se marcan varios de una vez y el pie lleva la cuenta a cada clic: «hay que quitar "
     "unos 38,2 g» → «aún sobrarían 27,6» → … → «con eso ya cuadra». La misma comida pasa de "
     "CINCO preguntas a UNA, y acaba igual de cuadrada.\n\n"
     "Cada línea decía además «quitándolo aún sobrarían N», que es el número de quitar ESE "
     "SOLO: en cuanto marcabas dos, cada línea decía una cosa y el pie otra. La línea dice lo "
     "que pone el alimento y la cuenta la lleva el pie.",
     [("_cuadra_hasta_cuadrar/1b_marcados.png", 3.4,
       "Cinco marcados y el pie diciendo que ya cuadra.")]),

    ("Una favorita de tres comidas decía que tenía cinco",
     "No está en el documento: lo reportó un cliente el 31-08",
     "«Cuando guardo una dieta para poder usarlo otro día, aparte de descuadrarme las "
     "cantidades, me sale esto» — y en la lista, su día de tres comidas aparecía con seis.",
     "El contador sumaba TODAS las claves con alimentos, y el intra y el post son claves. Un "
     "día de tres comidas con peri salía como «5 comidas»; y si además arrastraba la Comida 4 "
     "fantasma de antes del 29-08, como «6».\n\n"
     "Medido en producción de paso: de todas las favoritas de día de la base, solo UNA lleva "
     "más comidas montadas que su propio número, así que lo de la Comida 4 fantasma está "
     "prácticamente limpio desde aquel arreglo. Lo que seguía vivo era contar el peri.",
     [("_favorita_contador/1_tres_mas_peri.png", 3.4,
       "La misma favorita: «3 comidas + Intra y Post».")]),

    ("El filtro «Genérico» sacaba igual todas las marcas",
     "No está en el documento: lo reportó un cliente el 31-08",
     "«Cuando empiezo a hacer una dieta y selecciono alimentos genéricos, me salen también "
     "otros alimentos y todas sus marcas.»",
     "La regla del chip ya era buena — genérico = sin ficha de producto —. Lo que no "
     "funcionaba es que la búsqueda por texto no mandaba los chips: se mandaba solo lo "
     "escrito, así que el filtro seguía pintado y encendido pero dejaba de existir. Un filtro "
     "que se ve encendido y no filtra es peor que no tenerlo, porque el cliente cree que la "
     "lista ya está filtrada. El mismo agujero por otra puerta: encender el chip con el texto "
     "ya escrito tampoco rehacía la lista.\n\n"
     "Y una segunda cosa que encontró Francisco: «pero hay marcas que no tienen url». Tenía "
     "razón. Había SEIS fichas de marca sin URL que se colaban en «Genérico» — 7 Hermanos, "
     "Sol Natural, Nutrisport (dos), un Hacendado y Esgir —, ya marcadas a mano. No se les "
     "inventa una URL: una URL falsa acaba siendo un enlace que se le enseña a un cliente.\n\n"
     "Y hay una prueba que salta cuando aparezca otra, porque si no, dentro de dos meses "
     "volvemos a tener marcas coladas y nadie se entera.",
     [("_generico/1_sin_filtro.png", 3.0,
       "Sin filtro: Aldelís, Aldi, Frial, La Selva, Lidl, Hacendado…"),
      ("_generico/2_con_generico.png", 3.0,
       "Con «Genérico» encendido, y el texto ya escrito: ni una marca.")]),
]


def main():
    doc = Document()
    est = doc.styles["Normal"]
    est.font.name = "Segoe UI"
    est.font.size = Pt(10)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    # ── Portada ──
    p = doc.add_paragraph()
    r = p.add_run("12EN12 · 31 DE AGOSTO DE 2026")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = NARANJA

    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(6)
    r = t.add_run("«El día», punto por punto")
    r.bold = True
    r.font.size = Pt(24)

    parrafo(doc,
            "El documento se queda como está. Esto es lo otro: por cada cosa que pide, qué "
            "pedía, qué hace hoy el sistema y una captura de la app de verdad.",
            tam=10.5, color=GRIS, despues=6)
    parrafo(doc,
            "El bloque F no sale del documento: son los fallos que reportaron los clientes y "
            "Francisco usando la app ese mismo día, con la nota de voz de Jesús sobre el "
            "botón de cuadrar. Van aquí para que todo lo del 31 esté en un sitio.",
            tam=9.5, color=GRIS, despues=6)
    parrafo(doc,
            "Las capturas están sacadas con la app corriendo y con una cuenta de prueba, no "
            "son maquetas. Las que dependen de la hora se sacan forzando la respuesta del "
            "servidor, que es lo mismo que el navegador recibe a esa hora; la regla de las "
            "horas la prueban 30 casos automáticos con los bordes al minuto (14:59 contra "
            "15:00, y 16:59 contra 17:00).",
            tam=9, color=GRIS, despues=14)

    for punto in PUNTOS:
        nombre, donde, pedia, hace, fotos = punto[0], punto[1], punto[2], punto[3], punto[4]
        if donde is None:      # es un separador de bloque
            titulo(doc, nombre, tam=13, color=NARANJA, espacio_antes=22)
            continue

        titulo(doc, nombre, tam=12, espacio_antes=18)
        parrafo(doc, f"En el documento: {donde}", tam=8, color=GRIS, despues=6)

        if pedia and pedia != "—":
            etiqueta(doc, "lo que pide el documento")
            cita(doc, pedia)

        etiqueta(doc, "cómo está el sistema ahora")
        for trozo in (hace or "").split("\n\n"):
            parrafo(doc, trozo, tam=10)

        # Hay puntos sin captura posible (los que se ven en los números, no en la pantalla):
        # ahí no se pone el rótulo, que si no queda un «la prueba» sin nada debajo.
        if fotos:
            etiqueta(doc, "la prueba")
            for fichero, ancho, pie in fotos:
                captura(doc, fichero, ancho, pie)

    # ── Cierre ──
    titulo(doc, "Lo que no se ve en una captura", tam=13, color=NARANJA, espacio_antes=24)
    parrafo(doc,
            "Dos cosas del documento no tienen pantalla que enseñar, y van aquí para que no "
            "parezcan olvidadas.", tam=10)
    parrafo(doc,
            "«Las notificaciones del móvil»: no hay interruptor porque no hay "
            "notificaciones. «Un interruptor que no hace nada enseña que la configuración "
            "miente.» Cuando existan, se añade.", tam=10)
    parrafo(doc,
            "«El aviso que falta: cuando tú le contestas»: el documento dice que no está "
            "diseñado en ninguna parte, y sí lo está. Los cinco casos que nombra ya avisan "
            "hoy, cada uno al guardarse: «Ya tienes tus macros nuevos» (y «Este mes no te "
            "toco nada» si no le tocas nada), «Ya tienes tu rutina», «Tu informe está listo», "
            "«Hemos comentado tu reporte» y el aviso del chat. Y uno más que el documento no "
            "nombra: «Hemos comentado tu check-in».", tam=10)
    parrafo(doc,
            "Lo que sí faltaba era el otro lado: al mandar el reporte se le promete una fecha "
            "—«Antes del viernes tienes tus ajustes nuevos»— y si nadie le contesta no pasa "
            "nada. Ahora salta un aviso al equipo el mismo día prometido y a las 10:00, con "
            "media jornada por delante: el objetivo es que la promesa se cumpla, no dejar "
            "constancia de que se rompió. No se le avisa a él: sería la app anunciándole que "
            "le hemos fallado, y eso no le da sus ajustes.", tam=10)

    doc.save(SALIDA)
    print("Word ->", SALIDA)


if __name__ == "__main__":
    main()
