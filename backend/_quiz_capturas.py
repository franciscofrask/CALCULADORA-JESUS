# -*- coding: utf-8 -*-
"""DEV: recorre el Quiz Inicial completo como cliente demo, captura cada pantalla
y genera un Word en el Escritorio con todas las preguntas.
Requiere: frontend en :3000, cliente demo con questionnaire_completed=False y
calculadora=personalizado (ver _quiz_reset_demo.py --reset)."""
import os, sys, time
sys.stdout.reconfigure(encoding="utf-8")
from playwright.sync_api import sync_playwright

BASE = "http://localhost:3000"
EMAIL = "clientedemo@test.com"
PW = "demo123"
SHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_quiz_shots")
DOCX = r"C:\Users\Administrador\Desktop\Quiz Inicial - preguntas.docx"
os.makedirs(SHOTS, exist_ok=True)

captures = []  # (path, caption)

def run():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_context(viewport={"width": 1300, "height": 950},
                                   device_scale_factor=2).new_page()
        page.set_default_timeout(15000)

        n = {"i": 0}
        def shot(caption, wait_title=None, pre_wait=0.5):
            if wait_title:
                page.get_by_text(wait_title, exact=False).first.wait_for(state="visible")
            time.sleep(pre_wait)
            n["i"] += 1
            path = os.path.join(SHOTS, f"{n['i']:02d}.png")
            page.screenshot(path=path, full_page=True)
            captures.append((path, f"{n['i']:02d}. {caption}"))
            print(f"  captura {n['i']:02d}: {caption}")

        def click_btn(text):
            page.get_by_role("button", name=text, exact=False).first.click()

        def click_opt(text):
            page.locator("button", has_text=text).first.click()

        # ---- login ----
        page.goto(f"{BASE}/auth")
        page.get_by_placeholder("Email").fill(EMAIL)
        page.get_by_placeholder("Contraseña").fill(PW)
        page.get_by_role("button", name="Entrar").click()
        page.wait_for_url("**/questionnaire", timeout=20000)
        # ocultar toasts para capturas limpias
        page.add_style_tag(content="[data-sonner-toaster]{display:none!important}")
        time.sleep(1)

        # ===== NIVEL 0 =====
        shot("Portada - Quiz Inicial", wait_title="QUIZ INICIAL")
        click_btn("Empezar")

        shot("Nombre y apellidos", wait_title="Nombre y apellidos")
        click_btn("OK")

        shot("Número de teléfono", wait_title="Número de teléfono")
        page.locator("input").first.fill("600123456")
        click_btn("OK")

        shot("¿Cuál es tu sexo?", wait_title="¿Cuál es tu sexo?")
        click_opt("Hombre")

        shot("¿Cuál es tu objetivo?", wait_title="¿Cuál es tu objetivo?")
        click_opt("VOLUMEN")

        shot("¿Estás seguro? (confirmación objetivo)", wait_title="¿Estás seguro?")
        click_opt("Sí, lo tengo claro")

        shot("¿Cuánto pesas?", wait_title="¿Cuánto pesas?")
        page.locator("input").first.fill("80")
        click_btn("OK")

        shot("¿Cuál es tu porcentaje de grasa? (slider)", wait_title="porcentaje de grasa")
        click_btn("Continuar")

        shot("Afina tus macros (intro ajustes)", wait_title="Afina tus macros")
        click_btn("Vamos")

        shot("Actividad diaria fuera del gimnasio", wait_title="actividad diaria")
        click_opt("Normal: me muevo a diario")

        shot("¿Practicas otro deporte?", wait_title="otro deporte")
        click_opt("No")

        shot("Cuando te pasas comiendo, ¿engordas?", wait_title="¿engordas?")
        click_opt("Normal: si me paso")

        shot("¿Sigues una dieta ahora mismo?", wait_title="¿Sigues una dieta")
        click_opt("Sí, sé lo que como")
        # El paso 'dieta' es condicional; el auto-avance (closure con estado previo)
        # lo salta hasta 'final0'. Volvemos atrás para revelarlo ya con sigue_dieta=true.
        page.get_by_text("Y ya estaría", exact=False).first.wait_for(state="visible")
        page.get_by_role("button", name="Atrás").first.click()

        shot("Cuéntanos qué comes (dieta actual)", wait_title="Cuéntanos qué comes")
        page.get_by_placeholder("250").fill("250")
        page.get_by_placeholder("60").fill("60")
        page.locator("textarea").first.fill("Desayuno: avena y huevos. Comida: pollo con arroz y verduras. Cena: pescado con patata.")
        click_btn("OK")

        shot("Resumen antes de calcular", wait_title="Y ya estaría")
        click_btn("Calcular mis macros")

        shot("Resultado: tus macros", wait_title="Estos son tus macros", pre_wait=1.2)
        click_btn("Continuar")

        # ===== ONBOARDING =====
        shot("Ahora, tus gustos (intro)", wait_title="Ahora, tus gustos")
        click_btn("Vamos")

        shot("¿Cuántas comidas al día?", wait_title="¿Cuántas comidas")
        click_opt("4 comidas")

        shot("¿Cuántos días entrenas por semana?", wait_title="¿Cuántos días entrenas")
        click_opt("4 días")

        shot("¿Cuándo sueles entrenar? (momento)", wait_title="¿Cuándo sueles entrenar?")
        click_opt("Después de la comida 1")

        shot("Tus alimentos (preferencias)", wait_title="Configura tus preferencias")
        page.locator("[data-testid=save-preferences-btn]").click()

        shot("Momento mágico: comidas que puedes comer hoy",
             wait_title="comidas que puedes comer hoy", pre_wait=1.5)
        click_btn("Continuar con tu perfil")

        # ===== EL COMPLETO (planes con entrenador) =====
        # Las 21 pantallas del bloque 4 del documento del 18-08. Este recorrido es el de
        # alguien que ACABA DE HACER EL BÁSICO, así que aquí no vuelven a salir el biotipo, la
        # altura, la experiencia, los pesos ni las dietas de antes: ya las contestó.
        shot("Ahora, tu perfil completo (intro)", wait_title="tu perfil completo")
        click_btn("Seguir")

        shot("¿Alguna enfermedad o patología?", wait_title="enfermedad o patología")
        click_opt("No")

        shot("¿Tomas medicación?", wait_title="algún tipo de medicación")
        click_opt("No")

        shot("¿Tratamiento hormonal tipo TRT?", wait_title="tratamiento hormonal")
        click_opt("No")

        shot("¿Ayudas farmacológicas?", wait_title="ayudas farmacológicas")
        click_opt("No.")

        shot("¿Cuántas horas duermes?", wait_title="horas duermes")
        click_opt("Mínimo 7 y media")

        shot("¿Algo para dormir mejor?", wait_title="para dormir mejor")
        click_opt("No, duermo sin problema.")

        shot("¿Inconveniente con la suplementación?", wait_title="suplementación deportiva")
        click_opt("No. De hecho, suelo utilizarlos")

        shot("¿Qué suplementos tomas ahora?", wait_title="suplementos tomas ahora")
        page.locator("textarea").first.fill("Creatina y proteína de suero.")
        click_btn("OK")

        shot("¿Alguno contraindicado?", wait_title="contraindicado")
        page.locator("textarea").first.fill("no")
        click_btn("OK")

        shot("¿Entrenas ahora mismo?", wait_title="Entrenas ahora mismo")
        click_opt("Sí, voy mínimo 3 días")

        shot("¿Haces cardio?", wait_title="¿Haces cardio?")
        click_opt("Sí, además me gusta")

        shot("¿Arrastras alguna lesión?", wait_title="lesión o molestia")
        click_opt("No")

        shot("¿Con qué material cuentas?", wait_title="¿Con qué material cuentas")
        click_opt("Gimnasio completo")
        click_btn("OK")

        shot("Maquinaria que te falta", wait_title="En lo referente a maquinaria")
        page.locator("textarea").first.fill("no")
        click_btn("OK")

        shot("Tus fotos y tus medidas", wait_title="Tus fotos y tus medidas")
        click_btn("Ahora no")

        shot("Perfil completo (final)", wait_title="Perfil completo")
        # no pulsamos Enviar para no cerrar el flujo; ya tenemos todas las pantallas

        browser.close()

def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    t = doc.add_heading("Quiz Inicial - Todas las preguntas", level=0)
    sub = doc.add_paragraph("Cuestionario de onboarding (motor de macros v2): Nivel 0 + Onboarding + Nivel 1 (planes con coach). Capturas del flujo completo.")
    sub.runs[0].italic = True
    for path, caption in captures:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(4.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].bold = True
        cap.runs[0].font.size = Pt(10)
    doc.save(DOCX)
    print(f"\nWord generado: {DOCX}  ({len(captures)} pantallas)")

if __name__ == "__main__":
    run()
    build_docx()
