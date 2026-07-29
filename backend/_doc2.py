# -*- coding: utf-8 -*-
"""DEV: recorrido COMPLETO paso a paso de la app (flujos cliente + admin, con modales
y procesos) capturando cada estado, y genera un Word en el Escritorio.
Requiere frontend :3000 y backend :8000. Cliente demo y admin francisco."""
import os, sys, time, traceback
sys.stdout.reconfigure(encoding="utf-8")
from playwright.sync_api import sync_playwright

BASE = "http://localhost:3000"
DEMO = ("clientedemo@test.com", "demo123")
ADMIN = ("francisco@test.com", "demo123")
FICHA = "/admin/clients/20e8b4c4-98d1-43bc-ad74-06df9ddc8c94"
SHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_doc2_shots")
DOCX = r"C:\Users\Administrador\Desktop\12EN12 - Manual completo (flujos).docx"
os.makedirs(SHOTS, exist_ok=True)

CAPS = []   # dicts: part, flow, title, expl, path
_n = [0]

def make(page):
    def settle(t=1.0):
        try: page.wait_for_load_state("networkidle", timeout=5000)
        except Exception: pass
        try: page.add_style_tag(content="[data-sonner-toaster]{display:none!important}")
        except Exception: pass
        page.wait_for_timeout(int(t * 1000))

    def cap(part, flow, title, expl):
        _n[0] += 1
        path = os.path.join(SHOTS, f"{_n[0]:03d}.png")
        try:
            page.wait_for_timeout(350)
            dlg = ""
            try:
                if page.get_by_role("dialog").first.is_visible(timeout=400): dlg = " [modal]"
            except Exception: pass
            page.screenshot(path=path)
            CAPS.append({"part": part, "flow": flow, "title": title, "expl": expl, "path": path})
            print(f"  {_n[0]:03d}{dlg} [{flow}] {title}")
        except Exception as e:
            print(f"  ERROR cap {title}: {e}")

    def goto(path, t=1.4):
        try:
            page.goto(f"{BASE}{path}"); settle(t)
            return True
        except Exception as e:
            print(f"  ERROR goto {path}: {e}"); return False

    def _loc(sel):
        return page.locator(sel).first

    def ct(testid, t=6000):
        try:
            page.locator(f'[data-testid="{testid}"]').first.click(timeout=t); page.wait_for_timeout(500); return True
        except Exception: return False

    def cx(text, t=5000):
        try:
            page.get_by_role("button", name=text, exact=False).first.click(timeout=t); page.wait_for_timeout(500); return True
        except Exception:
            try:
                page.locator(f'button:has-text("{text}")').first.click(timeout=2500); page.wait_for_timeout(500); return True
            except Exception: return False

    def ctext(text, t=4000):
        try:
            page.get_by_text(text, exact=False).first.click(timeout=t); page.wait_for_timeout(500); return True
        except Exception: return False

    def ft(testid, val, t=4000):
        try:
            el = page.locator(f'[data-testid="{testid}"]').first
            el.fill(str(val), timeout=t); page.wait_for_timeout(200); return True
        except Exception: return False

    def fp(placeholder, val, idx=0):
        try:
            page.get_by_placeholder(placeholder, exact=False).nth(idx).fill(str(val), timeout=3000); return True
        except Exception: return False

    def st(testid, val):
        try:
            page.locator(f'[data-testid="{testid}"]').first.select_option(str(val), timeout=3000); page.wait_for_timeout(400); return True
        except Exception:
            try:
                page.locator(f'[data-testid="{testid}"]').first.select_option(label=str(val), timeout=2000); page.wait_for_timeout(400); return True
            except Exception: return False

    def ctf(testid):
        try:
            el = page.locator(f'[data-testid="{testid}"]').first
            el.scroll_into_view_if_needed(timeout=3000); el.click(force=True, timeout=4000); page.wait_for_timeout(600); return True
        except Exception: return False

    def cvis(text):
        try:
            page.locator(f'button:has-text("{text}"):visible').first.click(force=True, timeout=4000); page.wait_for_timeout(600); return True
        except Exception: return False

    def esc():
        try: page.keyboard.press("Escape"); page.wait_for_timeout(500)
        except Exception: pass

    def login(email, pw):
        page.goto(f"{BASE}/auth")
        try: page.evaluate("() => localStorage.clear()")
        except Exception: pass
        page.goto(f"{BASE}/auth"); page.wait_for_timeout(800)
        try:
            page.get_by_placeholder("Email").fill(email)
            page.get_by_placeholder("Contraseña").fill(pw)
            page.get_by_role("button", name="Entrar").click()
            page.wait_for_timeout(2800)
        except Exception as e:
            print("  ERROR login:", e)

    return dict(settle=settle, cap=cap, goto=goto, ct=ct, cx=cx, ctext=ctext,
                ft=ft, fp=fp, st=st, esc=esc, login=login, page=page, ctf=ctf, cvis=cvis)


def run():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        ctx = browser.new_context(viewport={"width": 1400, "height": 1000}, device_scale_factor=2)
        page = ctx.new_page(); page.set_default_timeout(15000)
        H = make(page)
        g, cap, ct, cx, ctext, ft, fp, st, esc = (H["goto"], H["cap"], H["ct"], H["cx"],
            H["ctext"], H["ft"], H["fp"], H["st"], H["esc"])
        ctf, cvis = H["ctf"], H["cvis"]

        def flow(fn):
            try: fn()
            except Exception as e:
                print(f"  FLOW ERROR {fn.__name__}: {e}"); traceback.print_exc()

        # ============================ CLIENTE ============================
        H["login"](*DEMO)

        def f_nutricion():
            P, F = "App del cliente", "Nutrición (montar la dieta)"
            g("/dashboard/nutrition", 2)
            cap(P, F, "Modal de intro 'Prepara tu día en 3 pasos'",
                "Al entrar en Nutrición aparece un onboarding breve que resume el flujo: elegir el día, preparar las comidas y seguir los macros. Se cierra con 'Empezar'.")
            ct("start-intro-btn") or ct("close-intro-btn") or cx("Empezar")
            H["settle"](0.8)
            cap(P, F, "Pantalla principal de Nutrición",
                "El constructor de dietas. Arriba, acciones (PDF, Copiar, Favoritas, Preferencias); debajo, el resumen de macros del día, la configuración del día y la lista de comidas.")
            ct("tipo-dia-entrenamiento")
            st("num-comidas-select", "4"); st("momento-entreno-select", "2"); st("peri-select", "intra_post")
            cap(P, F, "Configuración del día",
                "Se define el contexto: día de entrenamiento o descanso, número de comidas, cuándo se entrena y la opción de perientreno (intra/post). Esto ajusta el reparto de macros y el orden de las comidas.")
            cap(P, F, "Lista de comidas del día",
                "Cada comida (C1..C4, Intra, Post) tiene botones para construirla ('Lo hago yo'), sugerir un menú o repetir de otro día, con su estado (vacía, cuadrada, falta) y barras de macros.")
            # abrir BuildMealModal en C1 (boton 'Lo hago yo')
            ct("build-meal-C1")
            H["settle"](1.0)
            cap(P, F, "Modal 'Construir comida' (paso 1: proteína)",
                "Asistente para montar una comida. Modo Automático o Manual. Primero la proteína: al elegir una categoría (aves, huevos, pescados...) la calculadora sugiere el alimento y la cantidad, mostrando los macros objetivo de la comida.")
            (ctext("Aves") or ctext("Huevos") or ctext("Pollo") or ctext("Pescados"))
            H["settle"](0.6)
            cap(P, F, "Selección de categoría y lista de alimentos",
                "Elegida la categoría, aparece la lista de alimentos con su cantidad sugerida y macros. Al pulsar uno se añade a la comida y las barras de progreso se actualizan (pasa a hidratos y luego a grasas para cuadrar).")
            esc()
            # Sugiereme un menu (LibraryMenusModal)
            ct("menu-options-C1") or cx("Sugiéreme un menú")
            H["settle"](1.0)
            cap(P, F, "Modal 'Sugiéreme un menú' (biblioteca de menús reales)",
                "Busca en la biblioteca de menús reales de otros clientes que ya cuadran con el objetivo de esa comida. Se puede ajustar el margen, ordenar por más cuadrado o más usado, ver método/reales y filtrar por alimento.")
            esc()
            # Favoritas (Dialog)
            ct("open-favorites-btn") or cx("Favoritas")
            H["settle"](0.6)
            cap(P, F, "Modal 'Dietas favoritas'",
                "Guarda la dieta del día como plantilla reutilizable con nombre, y permite aplicar una favorita (adaptándola si cambia el tipo de día entreno/descanso).")
            esc()
            # Calendario (Dialog)
            ctf("open-calendar-btn")
            H["settle"](0.6)
            cap(P, F, "Modal 'Calendario' de dietas",
                "Vista mensual con el estado de cada día: cuadrado (verde), sin cuadrar (naranja), sin dieta (gris) y cambios de macros (azul). Al pulsar un día se navega a esa fecha.")
            esc()
            # Copiar dieta (Dialog)
            cvis("Copiar")
            H["settle"](0.6)
            cap(P, F, "Modal 'Copiar dieta'",
                "Copia la dieta completa del día actual a otra fecha futura, para no montarla de cero.")
            esc()
            # Preferencias (overlay a pantalla completa: va al final del flujo)
            ct("open-preferences-btn") or cx("Preferencias")
            H["settle"](0.7)
            cap(P, F, "Preferencias de alimentos (Me gusta)",
                "El cliente marca las categorías de alimentos que le gustan (mínimo 3). Estas preferencias afinan las sugerencias de la calculadora.")
            (ctext("Evito") or cx("Evitar"))
            cap(P, F, "Preferencias de alimentos (Evitar)",
                "Pestaña de alimentos a evitar: categorías completas y palabras clave (alergias/intolerancias o lo que no le gusta). Nunca aparecerán en las sugerencias.")

        def f_reportes():
            P, F = "App del cliente", "Reportes (seguimiento)"
            g("/dashboard/reports", 1.6)
            cap(P, F, "Formulario de reporte (vacío)",
                "El cliente registra su seguimiento: peso, medidas (pecho, cintura, cadera, brazo, muslo) y cumplimientos/sensaciones con deslizadores.")
            ft("weight-input", "75.5")
            fp("--", "98.5", 0); fp("--", "82", 1); fp("--", "95", 2); fp("--", "32.5", 3); fp("--", "56", 4)
            ft("notes-textarea", "Semana muy productiva, buena adherencia en entreno y nutrición.")
            cap(P, F, "Formulario relleno",
                "Con el peso, las medidas y las notas cargadas, el cliente pulsa 'Enviar reporte' y sus datos quedan registrados para el coach.")
            ct("submit-report-btn")
            H["settle"](1.6)
            cap(P, F, "Historial / evolución tras enviar",
                "Al enviar, la app pasa al historial y muestra el reporte guardado. En la pestaña de evolución se ve la gráfica de la curva de peso.")

        def f_checkins():
            P, F = "App del cliente", "Check-ins"
            g("/dashboard/checkins", 1.4)
            cap(P, F, "Check-in diario rápido",
                "Registro de 10 segundos: estado de ánimo con caras, energía, si entrenó y si siguió la dieta, más una nota opcional.")
            (ctext("Genial") or ctext("Bien"))
            cap(P, F, "Check-in con ánimo seleccionado",
                "El cliente marca su estado y envía; también puede desplegar el check-in semanal (peso y adherencias) y ver el historial de todos sus registros.")

        def f_mensajes():
            P, F = "App del cliente", "Mensajes"
            g("/dashboard/messages", 1.4)
            cap(P, F, "Chat con el entrenador",
                "Conversación con el coach o soporte: mensajes propios a la derecha (naranja) y del entrenador a la izquierda, con marcas de leído.")
            ft("message-input", "¿Qué recomiendas para la próxima semana? Quiero enfocar en volumen.")
            cap(P, F, "Escribiendo un mensaje",
                "El cliente escribe y pulsa enviar; el mensaje aparece al instante en la conversación con estado de lectura.")

        def f_perfil():
            P, F = "App del cliente", "Perfil y seguridad"
            g("/dashboard/profile", 1.4)
            cap(P, F, "Mi perfil",
                "Datos del cliente y su plan. Botón de editar para cambiar nombre y teléfono, y acceso a cambiar la contraseña.")
            ct("edit-profile-btn")
            cap(P, F, "Modo edición de perfil",
                "Formulario para actualizar nombre y teléfono, con guardar/cancelar.")
            esc()
            ct("edit-profile-btn")  # reabrir por si el esc cerró algo (idempotente)
            # modal cambiar contraseña
            (ctext("Cambiar contraseña") or cx("Cambiar contraseña"))
            H["settle"](0.6)
            cap(P, F, "Modal 'Cambiar contraseña'",
                "Cambio de contraseña con contraseña actual, nueva y confirmación (mínimo 8 caracteres).")
            esc()

        def f_chatbot():
            P, F = "App del cliente", "Asistente IA"
            g("/dashboard/chatbot", 1.4)
            cap(P, F, "Asistente de nutrición (inicio)",
                "Un asistente conversacional que ayuda a montar la dieta hablando. Se empieza con 'Empezar'.")
            cx("Empezar"); H["settle"](1.5)
            (cx("Hoy") or ctext("Hoy")); H["settle"](1.2)
            (cx("Día de Entrenamiento") or ctext("entrenamiento")); H["settle"](1.2)
            cap(P, F, "Conversación de configuración",
                "El asistente pregunta día, tipo de día, número de comidas y perientreno, y el cliente responde con botones. Luego pide los alimentos comida a comida.")

        def f_suplementos():
            P, F = "App del cliente", "Suplementación"
            g("/dashboard/supplements", 1.4)
            cap(P, F, "Mi suplementación",
                "Protocolo del cliente (actual y siguiente, con su fecha). Cada suplemento indica cuándo tomarlo, cuánto, observaciones y enlaces de compra.")

        def f_foods():
            P, F = "App del cliente", "Buscador de alimentos"
            g("/dashboard/foods", 1.4)
            fp("Texto en el alimento", "pollo")
            H["settle"](0.8)
            cap(P, F, "Buscador de alimentos",
                "Búsqueda por nombre con filtros por categoría. Cada resultado muestra sus macros y categorías; los de marca llevan enlace.")
            (cx("Sugerir alimento") or ctext("Sugerir alimento"))
            H["settle"](0.8)
            cap(P, F, "Modal 'Sugerir alimento'",
                "El cliente propone un alimento nuevo: nombre, si va por unidad, macros y fotos del producto y de la tabla nutricional. El equipo lo revisa.")
            esc()

        def f_macrocalc():
            P, F = "App del cliente", "Calculadora de macros"
            g("/dashboard/macro-calculator", 1.4)
            cap(P, F, "Calculadora / editor de macros",
                "Herramienta avanzada: a la izquierda calcula macros a partir de peso, % graso, sexo y objetivo; a la derecha el editor manual por día (entreno/descanso/peri) con fecha de vigencia.")
            fp("80", "82.5", 0); fp("20", "18.5", 0)
            (cx("Calcular"))
            H["settle"](1.2)
            cap(P, F, "Resultado del cálculo",
                "Tras calcular, muestra los macros para día de entreno, perientreno y descanso, con opción de cargarlos en el editor y guardarlos desde una fecha.")

        for fn in (f_nutricion, f_reportes, f_checkins, f_mensajes, f_perfil,
                   f_chatbot, f_suplementos, f_foods, f_macrocalc):
            flow(fn)

        # ============================ ADMIN ============================
        H["login"](*ADMIN)

        def f_admin_dash():
            P, F = "Panel admin/entrenador", "Dashboard admin"
            g("/admin", 1.6)
            cap(P, F, "Panel de control",
                "Estado del negocio en tiempo real: clientes totales/activos/en riesgo, bajas, MRR y distribución por plan, más macros por revisar y próximos cobros.")
            ct("notif-bell")
            cap(P, F, "Campanita de notificaciones",
                "Desplegable con leads nuevos sin gestionar y mensajes sin leer, para saltar directamente a ellos.")
            esc()

        def f_ficha():
            P, F = "Panel admin/entrenador", "Ficha del cliente (acciones)"
            g(FICHA, 1.8)
            # Macros
            ct("tab-macros"); H["settle"](0.6)
            cap(P, F, "Pestaña Macros",
                "Historial de macros del cliente. Desde 'Cambiar macros' se registra un nuevo ajuste.")
            ct("change-macros-btn"); H["settle"](0.6)
            ft("macro-input-tp", "160")
            cap(P, F, "Modal 'Cambiar macros'",
                "Se fijan los macros de día de entreno, descanso y perientreno, la fecha de vigencia y el motivo del cambio. Guardar crea una nueva entrada en el historial.")
            esc()
            # Entreno
            ct("tab-entrenamiento"); H["settle"](0.6)
            ft("routine-instructions", "Rutina full-body 3 días, foco en espalda y pecho, sin peso muerto.")
            cap(P, F, "Pestaña Entreno (generar rutina con IA)",
                "Muestra equipamiento/lesiones y la rutina vigente. El coach da instrucciones y la IA genera una rutina que puede revisar y guardar.")
            # Menus
            ct("tab-menus"); H["settle"](0.6)
            ft("menufinder-P", "160"); ft("menufinder-H", "250"); ft("menufinder-G", "70")
            cap(P, F, "Pestaña Menús (buscador para el cliente)",
                "El coach busca menús reales que cuadren con unos macros objetivo, y los copia o los envía directamente al chat del cliente.")
            ct("menufinder-buscar"); H["settle"](1.6)
            cap(P, F, "Buscador de menús en acción",
                "Al buscar, la app devuelve menús reales que cuadran con esos macros; cada uno se copia o se envía al chat del cliente. (En este entorno de pruebas la biblioteca es reducida, por eso puede salir vacío; en producción hay cientos de miles.)")
            # Suplementos
            ct("tab-suplementos"); H["settle"](0.6)
            cap(P, F, "Pestaña Suplementos",
                "Protocolo actual y siguiente del cliente, con auto-sugerencia, añadir desde el catálogo y guardar. Debajo, la suplementación importada de Calma.")
            # Seguimiento
            ct("tab-seguimiento"); H["settle"](0.8)
            cap(P, F, "Pestaña Seguimiento",
                "Evolución de peso, check-ins y reportes con feedback del coach, más los reportes mensuales y las fotos de progreso importadas de Calma.")

        def f_leads():
            P, F = "Panel admin/entrenador", "Leads (CRM)"
            g("/admin/leads", 1.6)
            cap(P, F, "Vista Kanban del CRM",
                "Los prospectos en columnas por estado (nuevo, contactado, llamada agendada, propuesta enviada, convertido, descartado). Se arrastran entre columnas.")
            ct("add-lead-btn"); H["settle"](0.6)
            ft("new-lead-name", "Sergio Molina"); ft("new-lead-email", "sergio.molina@gmail.com")
            cap(P, F, "Modal 'Nuevo lead'",
                "Alta manual de un prospecto: nombre, email, teléfono, origen, responsable y notas.")
            esc()
            # abrir detalle de un lead (tarjeta kanban)
            try:
                page.locator('[data-testid^="kanban-card-"]').first.click(timeout=4000); page.wait_for_timeout(800)
            except Exception:
                ct("view-tabla")
                try: page.locator('[data-testid^="table-row-"]').first.click(timeout=3000); page.wait_for_timeout(800)
                except Exception: pass
            cap(P, F, "Detalle de un lead",
                "Ficha del prospecto: estado (botones para moverlo), responsable, próxima acción, notas e historial de interacciones. Botones para convertir o descartar.")
            ct("convert-btn"); H["settle"](0.6)
            cap(P, F, "Convertir lead en cliente",
                "Se elige plan y entrenador; al confirmar se crea el usuario cliente y se muestra una contraseña temporal y un mensaje de bienvenida listo para WhatsApp.")
            esc(); esc()
            ct("view-metricas"); H["settle"](1.0)
            cap(P, F, "Métricas del CRM",
                "Indicadores de conversión: leads totales, convertidos, tasa de conversión, días hasta convertir, y gráficos de embudo, origen y motivos de descarte.")

        def f_planes():
            P, F = "Panel admin/entrenador", "Planes"
            g("/admin/planes", 1.6)
            cap(P, F, "Catálogo de planes",
                "Todos los planes por estado (activos, legacy, especiales). Cada tarjeta muestra ciclo, precio, calculadora, rutina, reportes y suplementación.")
            try:
                page.locator('[data-testid^="edit-plan-"]').first.click(timeout=4000); page.wait_for_timeout(800)
                cap(P, F, "Modal 'Editar plan'",
                    "Permite ajustar los parámetros del plan: nombre, ciclo y semanas, precio, tipo de calculadora, rutina, cadencia de reportes y suplementación, con vista previa de lo que incluye.")
                esc()
            except Exception: pass

        def f_usuarios():
            P, F = "Panel admin/entrenador", "Usuarios (equipo)"
            g("/admin/usuarios", 1.6)
            cap(P, F, "Gestión del equipo",
                "Tabla de usuarios (admin y entrenadores) con filtros. Acciones por fila: editar, restablecer contraseña y dar de baja/reactivar.")
            ct("users-tab-actividad"); H["settle"](0.6)
            cap(P, F, "Pestaña Actividad (auditoría)",
                "Registro de auditoría: qué se cambió (macros, coach, rol, contraseña), quién lo hizo y cuándo.")
            ct("users-tab-equipo")

        def f_menus_admin():
            P, F = "Panel admin/entrenador", "Menús (plantillas)"
            g("/admin/menus", 1.6)
            cap(P, F, "Menús preestablecidos",
                "Plantillas de menús del recetario, filtrables por momento (desayuno, comida, merienda, cena).")
            cx("Nuevo menú"); H["settle"](0.6)
            cap(P, F, "Modal 'Nuevo menú'",
                "Creación de una plantilla: nombre, momento, etiquetas y los alimentos por rol (proteína/hidrato/grasa) con su proporción, buscándolos por nombre.")
            esc()

        def f_alimentos_admin():
            P, F = "Panel admin/entrenador", "Sugerencias de alimentos"
            g("/admin/alimentos", 1.6)
            cap(P, F, "Sugerencias de alimentos",
                "Alimentos propuestos por los clientes pendientes de revisión, con sus fotos y macros. El coach puede editar, aprobar o rechazar.")
            (cx("Editar") or ctext("Editar")); H["settle"](0.6)
            cap(P, F, "Modal de edición/aprobación",
                "Se revisa y corrige el alimento (nombre, por unidad, ración, macros, categoría) antes de aprobarlo e incorporarlo a la calculadora.")
            esc()

        def f_supcat():
            P, F = "Panel admin/entrenador", "Catálogo de suplementos"
            g("/admin/supplements-catalog", 1.6)
            cap(P, F, "Catálogo de suplementos",
                "Los suplementos que luego se asignan a los clientes, con su categoría, sexo, objetivo, timing y dosis.")
            cx("Nuevo"); H["settle"](0.6)
            cap(P, F, "Modal 'Nuevo suplemento'",
                "Alta de un suplemento: título, sexo/categoría/objetivo, cuándo y cuánto, observaciones, imagen y enlaces de compra.")
            esc()

        def f_admin_msg():
            P, F = "Panel admin/entrenador", "Mensajes (equipo)"
            g("/admin/messages", 1.6)
            cap(P, F, "Bandeja de mensajes",
                "A la izquierda las conversaciones con clientes (con no leídos) y a la derecha el chat abierto para responder.")
            try:
                page.locator('[data-testid^="conv-"]').first.click(timeout=4000); page.wait_for_timeout(900)
                cap(P, F, "Conversación abierta",
                    "El chat con un cliente concreto: mensajes propios a la derecha y del cliente a la izquierda, con campo para responder.")
            except Exception: pass

        def f_routines_admin():
            P, F = "Panel admin/entrenador", "Estado de rutinas"
            g("/admin/routines", 1.6)
            cap(P, F, "Estado de rutinas",
                "Vista general de qué clientes tienen rutina y cuáles no, con filtro 'solo sin rutina'. Cada fila lleva a la ficha del cliente (pestaña Entreno).")

        for fn in (f_admin_dash, f_ficha, f_leads, f_planes, f_usuarios,
                   f_menus_admin, f_alimentos_admin, f_supcat, f_admin_msg, f_routines_admin):
            flow(fn)

        browser.close()


def build():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.add_heading("12EN12 · Manual completo de la app (flujos)", level=0)
    intro = doc.add_paragraph("Recorrido paso a paso de la aplicacion siguiendo los flujos reales de uso: primero el circuito del CLIENTE (montar la dieta, reportes, check-ins, mensajes, perfil, asistente, suplementos, buscador y calculadora), con sus procesos y modales; despues el circuito del ADMIN/ENTRENADOR (dashboard, ficha del cliente con sus acciones, CRM de leads, planes, equipo, menus, alimentos, suplementos, mensajes y rutinas). Cada captura explica que se ve y que se hace en ese paso.")
    intro.runs[0].italic = True
    last_part = last_flow = None
    for c in CAPS:
        if c["part"] != last_part:
            doc.add_heading(c["part"], level=1); last_part = c["part"]; last_flow = None
        if c["flow"] != last_flow:
            doc.add_heading(c["flow"], level=2); last_flow = c["flow"]
        h = doc.add_paragraph(); r = h.add_run(c["title"]); r.bold = True; r.font.size = Pt(11)
        if c["path"] and os.path.exists(c["path"]):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(c["path"], width=Inches(6.2))
        doc.add_paragraph(c["expl"])
    doc.save(DOCX)
    print(f"\nWord generado: {DOCX}  ({len(CAPS)} capturas)")


if __name__ == "__main__":
    run()
    build()
