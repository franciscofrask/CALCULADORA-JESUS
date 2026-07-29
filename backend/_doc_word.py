# -*- coding: utf-8 -*-
"""DEV: recorre toda la app (desde despues del quiz) como cliente demo y como admin,
captura cada pantalla y genera un Word en el Escritorio explicando cada una.
Requiere frontend en :3000 y backend en :8000. Cliente demo con questionnaire_completed=True."""
import os, sys, time
sys.stdout.reconfigure(encoding="utf-8")
from playwright.sync_api import sync_playwright

BASE = "http://localhost:3000"
DEMO = ("clientedemo@test.com", "demo123")
ADMIN = ("francisco@test.com", "demo123")
FICHA_ID = "20e8b4c4-98d1-43bc-ad74-06df9ddc8c94"  # arcochat (datos ricos)
SHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_doc_shots")
DOCX = r"C:\Users\Administrador\Desktop\12EN12 - Guia de pantallas.docx"
os.makedirs(SHOTS, exist_ok=True)

# (seccion, titulo, rol, ruta|None, tab|None, explicacion)
FICHA = f"/admin/clients/{FICHA_ID}"
ENTRIES = [
    ("Después del quiz", "Bienvenida (post-quiz)", "client", "/welcome", None,
     "Es lo primero que ve el cliente al terminar el Quiz Inicial. Le muestra sus macros ya calculados (proteina, hidratos y grasa, con las kcal) y le da la bienvenida por su nombre. Desde aqui elige entre 'Empezar recorrido guiado' (un tour por la app) o 'Explorar por mi cuenta'."),
    ("Después del quiz", "Selección de plan (onboarding)", "client", "/onboarding", None,
     "Catalogo de planes contratables. Cada tarjeta muestra nombre, precio, descripcion y lo que incluye. Al pulsar 'Ir a pagar' abre el checkout de Stripe; al volver del pago, el perfil se actualiza y entra al panel. Un cliente ya activo (como el demo) puede saltarlo."),

    ("App del cliente", "Panel de inicio (dashboard)", "client", "/dashboard", None,
     "La home del cliente. Cabecera con su plan y accesos. Incluye un checklist de primeros pasos con barra de progreso, los macros del dia con anillos de cumplimiento, y tarjetas de acceso rapido a Rutina, Nutricion, Reportes, Mensajes, Perfil, Asistente, Suplementos, Calculadora y Alimentos. Campanita de notificaciones."),
    ("App del cliente", "Mi rutina (oculta temporalmente)", "note", None, None,
     "La pantalla de rutina del cliente (vista semanal con el selector de dia, los ejercicios de cada dia con series/repeticiones y el historial de rutinas) existe pero esta OCULTA temporalmente en la app hasta completar la funcionalidad, por eso no aparece en el menu lateral. La gestion de rutinas sigue disponible para el coach en el panel de admin (Estado de rutinas y la pestana Entreno de la ficha del cliente)."),
    ("App del cliente", "Mi nutrición (constructor de dietas)", "client", "/dashboard/nutrition", None,
     "El nucleo de la app: monta la dieta del dia paso a paso. Se elige fecha, tipo de dia (entreno/descanso), numero de comidas y el momento del entreno para colocar el perientreno. Cada comida se construye por bloques (proteina, acompanamiento, peri/post) y muestra los macros frente al objetivo. Incluye preferencias, copiar dieta, biblioteca de menus, favoritas y calendario."),
    ("App del cliente", "Evolución y reportes", "client", "/dashboard/reports", None,
     "Donde el cliente envia su reporte de seguimiento: peso, medidas (pecho, cintura, cadera, brazo, muslo), cumplimiento de entreno y nutricion, y sensaciones (sueno, energia, estres) con notas. La pestana de historial muestra la grafica de peso y los reportes anteriores."),
    ("App del cliente", "Check-ins diarios", "client", "/dashboard/checkins", None,
     "Version ligera y rapida del seguimiento para el dia a dia: peso, estado de salud, cumplimiento de calorias/nutricion/entreno y estado de animo con caras. Guarda el historial de check-ins con su estado. Permite adjuntar foto."),
    ("App del cliente", "Mensajes con el entrenador", "client", "/dashboard/messages", None,
     "Chat directo con el entrenador o con soporte. Muestra la conversacion con burbujas (entrantes y propias), marcas de leido y hora. Campo para escribir y enviar; se actualiza solo cada pocos segundos."),
    ("App del cliente", "Mi perfil", "client", "/dashboard/profile", None,
     "Datos del cliente (nombre, email, plan) con modo edicion para nombre y telefono. Permite cambiar la contrasena en un modal y ver lo que incluye su plan. Boton de cerrar sesion."),
    ("App del cliente", "Asistente (chatbot IA)", "client", "/dashboard/chatbot", None,
     "Asistente conversacional que ayuda a montar la dieta hablando: configura fecha, tipo de dia y comidas, y va proponiendo alimentos paso a paso. Mantiene la conversacion durante la sesion y puede volcar el resultado a Nutricion."),
    ("App del cliente", "Mi suplementación", "client", "/dashboard/supplements", None,
     "Protocolo de suplementos del cliente: bloque actual y bloque siguiente (con su fecha de inicio). Cada suplemento muestra imagen, cuando tomarlo, cuanto, observaciones y enlaces de compra. Incluye una nota personal del coach."),
    ("App del cliente", "Calculadora de macros", "client", "/dashboard/macro-calculator", None,
     "Editor manual y avanzado de macros por dia: proteina, hidratos y grasa para dia de entreno, de descanso y perientreno. Permite fijar una fecha, dejar una nota y copiar los macros a los proximos dias."),
    ("App del cliente", "Buscador de alimentos", "client", "/dashboard/foods", None,
     "Biblioteca de alimentos con buscador y filtros por categoria. Cada resultado muestra sus macros, categorias y cantidad minima; los de marca llevan enlace. El cliente puede proponer un alimento nuevo con el boton de sugerir."),

    ("Panel admin/trainer", "Panel de control (dashboard admin)", "admin", "/admin", None,
     "Home del staff con el estado del negocio en tiempo real: clientes totales/activos/en riesgo, bajas, MRR y distribucion por plan. Muestra los macros por revisar (dietas que no cuadran), los proximos cobros y la cadencia de reportes. Campanita con leads nuevos y mensajes sin leer."),
    ("Panel admin/trainer", "Lista de clientes", "admin", "/admin/clients", None,
     "Listado de clientes con buscador y filtros. Cada fila lleva a la ficha detallada del cliente. Es el punto de entrada para gestionar a cada persona."),
    ("Panel admin/trainer", "Ficha del cliente · Resumen", "admin", FICHA, "Resumen",
     "Cabecera de la ficha con datos del cliente y su plan, mas un resumen general (macros, peso, ultima actividad). Es la portada de la ficha, con las pestanas de todas las areas del cliente."),
    ("Panel admin/trainer", "Ficha del cliente · Macros", "admin", FICHA, "Macros",
     "Historial de macros del cliente por fecha, con la posibilidad de editar, repetir o eliminar entradas y de registrar un cambio con su motivo y fecha de vigencia."),
    ("Panel admin/trainer", "Ficha del cliente · Calculadora", "admin", FICHA, "Calculadora",
     "Calculadora de macros del coach para este cliente: recalcula sus numeros a partir de sus datos y respuestas, y permite aplicarlos al perfil."),
    ("Panel admin/trainer", "Ficha del cliente · Membresía", "admin", FICHA, "Membresía",
     "Gestion de la cuenta: rol, plan (incluido plan de cortesia), baja logica, precio, inicio y proximo cobro, e historial de pagos. Aqui tambien se ve el historial de membresias importado de Calma."),
    ("Panel admin/trainer", "Ficha del cliente · Cuestionario", "admin", FICHA, "Cuestionario",
     "Datos del cuestionario del cliente: objetivo, peso, sexo, % graso, edad y altura, mas el cuestionario inicial completo importado de Calma (telefono, profesion, medicacion, medidas, lesiones, dieta de ejemplo, etc.)."),
    ("Panel admin/trainer", "Ficha del cliente · Entreno", "admin", FICHA, "Entreno",
     "Equipamiento y lesiones del cliente, su rutina vigente y el generador de rutina con IA (instrucciones + generar, revisar y guardar)."),
    ("Panel admin/trainer", "Ficha del cliente · Nutrición", "admin", FICHA, "Nutrición",
     "Estadisticas de nutricion del cliente: total de dietas, top de alimentos y el listado de dietas por fecha para abrir y revisar cada una."),
    ("Panel admin/trainer", "Ficha del cliente · Menús", "admin", FICHA, "Menús",
     "Buscador de menus para el cliente: se indican los macros objetivo y el momento, y busca en el recetario y en la biblioteca de menus reales. Cada menu se puede copiar o enviar directamente al chat del cliente."),
    ("Panel admin/trainer", "Ficha del cliente · Suplementos", "admin", FICHA, "Suplementos",
     "Protocolo de suplementacion del cliente (actual y siguiente, con fecha), con auto-sugerencia, anadir desde catalogo y nota personal. Muestra tambien la suplementacion importada de Calma."),
    ("Panel admin/trainer", "Ficha del cliente · Seguimiento", "admin", FICHA, "Seguimiento",
     "Evolucion de peso, check-ins y reportes del cliente con feedback del coach. Incluye los reportes mensuales completos y las fotos de progreso importados de Calma."),
    ("Panel admin/trainer", "Catálogo de planes", "admin", "/admin/planes", None,
     "Gestion de todos los planes (activos, legacy, especiales). Cada tarjeta muestra ciclo, precio, tipo de calculadora, rutina, cadencia de reportes y suplementacion, y permite editar sus parametros."),
    ("Panel admin/trainer", "Gestión de equipo (usuarios)", "admin", "/admin/usuarios", None,
     "Solo para admin: gestion de usuarios del equipo (admin y entrenadores) y clientes. Editar datos y rol, dar de baja/reactivar, y registro de actividad (auditoria)."),
    ("Panel admin/trainer", "Leads (CRM)", "admin", "/admin/leads", None,
     "CRM de prospectos con vista Kanban por estado (nuevo, contactado, propuesta, convertido, descartado), lista y metricas. Permite crear leads, registrar la siguiente accion, convertir en cliente o descartar con motivo. Integrado con GHL."),
    ("Panel admin/trainer", "Centro de mensajes (staff)", "admin", "/admin/messages", None,
     "Bandeja del equipo: a la izquierda la lista de conversaciones con clientes (con no leidos) y a la derecha el chat abierto para responder. Se actualiza solo."),
    ("Panel admin/trainer", "Estado de rutinas", "admin", "/admin/routines", None,
     "Vista general de que clientes tienen rutina y cuales no, con sus dias de entreno y fecha. Filtra por 'solo sin rutina' y lleva a la ficha del cliente."),
    ("Panel admin/trainer", "Biblioteca de menús (plantillas)", "admin", "/admin/menus", None,
     "Gestion de las plantillas de menus del recetario: nombre, momento (desayuno/comida/merienda/cena), rango de kcal e ingredientes por rol. Crear, editar y eliminar."),
    ("Panel admin/trainer", "Sugerencias de alimentos", "admin", "/admin/alimentos", None,
     "Alimentos propuestos por los clientes pendientes de revision: fotos, macros y categoria, con acciones de aprobar, rechazar o editar antes de incorporarlos."),
    ("Panel admin/trainer", "Catálogo de suplementos", "admin", "/admin/supplements-catalog", None,
     "Gestor del catalogo de suplementos que luego se asignan a los clientes: titulo, imagen, enlaces, cuando/cuanto, sexo, categoria y objetivo. Crear, editar y desactivar."),
]

def run():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        ctx = browser.new_context(viewport={"width": 1400, "height": 1000}, device_scale_factor=2)
        page = ctx.new_page()
        page.set_default_timeout(20000)

        def login(email, pw):
            page.goto(f"{BASE}/auth")
            try:
                page.evaluate("() => localStorage.clear()")
            except Exception:
                pass
            page.goto(f"{BASE}/auth")
            page.get_by_placeholder("Email").fill(email)
            page.get_by_placeholder("Contraseña").fill(pw)
            page.get_by_role("button", name="Entrar").click()
            page.wait_for_timeout(2500)

        def settle():
            try:
                page.wait_for_load_state("networkidle", timeout=6000)
            except Exception:
                pass
            page.add_style_tag(content="[data-sonner-toaster]{display:none!important}")
            page.wait_for_timeout(1200)

        def shot(i):
            path = os.path.join(SHOTS, f"{i:02d}.png")
            page.screenshot(path=path)
            return path

        results = {}
        # ----- cliente -----
        login(*DEMO)
        for i, e in enumerate(ENTRIES):
            if e[2] != "client":
                continue
            try:
                page.goto(f"{BASE}{e[3]}"); settle()
                results[i] = shot(i); print(f"  [cliente] {i:02d} {e[1]}")
            except Exception as ex:
                print(f"  ERROR {i} {e[1]}: {ex}")
        # ----- admin -----
        login(*ADMIN)
        for i, e in enumerate(ENTRIES):
            if e[2] != "admin":
                continue
            try:
                if not page.url.endswith(e[3]):
                    page.goto(f"{BASE}{e[3]}"); settle()
                if e[4]:  # pestana de la ficha
                    page.get_by_role("tab", name=e[4], exact=False).first.click()
                    page.wait_for_timeout(1400)
                results[i] = shot(i); print(f"  [admin]   {i:02d} {e[1]}{' · '+e[4] if e[4] else ''}")
            except Exception as ex:
                print(f"  ERROR {i} {e[1]}: {ex}")
        browser.close()
        return results

def build(results):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.add_heading("12EN12 · Guía de pantallas de la app", level=0)
    intro = doc.add_paragraph("Recorrido por todas las pantallas de la aplicacion, en orden del circuito de uso, con una explicacion de cada una. Continua justo despues del Quiz Inicial (documentado aparte): empezamos en la pantalla de bienvenida y seguimos por toda la app del cliente y el panel de admin/entrenador.")
    intro.runs[0].italic = True
    last_section = None
    for i, e in enumerate(ENTRIES):
        section, titulo = e[0], e[1]
        if section != last_section:
            doc.add_heading(section, level=1)
            last_section = section
        doc.add_heading(titulo, level=2)
        path = results.get(i)
        if e[2] == "note":
            pass  # entrada solo texto (pantalla oculta / sin captura)
        elif path and os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(6.3))
        else:
            doc.add_paragraph("(captura no disponible)").runs[0].italic = True
        doc.add_paragraph(e[5])
    doc.save(DOCX)
    print(f"\nWord generado: {DOCX}  ({sum(1 for i in range(len(ENTRIES)) if results.get(i))}/{len(ENTRIES)} capturas)")

if __name__ == "__main__":
    r = run()
    build(r)
